"""增强工具集插件 — 提供文档解析、数据处理、文本分析、文件操作等全面工具

工具分类：
- 文档解析：docx_read, xlsx_read, csv_analyze, pdf_read
- 数据处理：json_query, yaml_parse, xml_parse, table_format
- 文本处理：regex_extract, text_diff, text_transform
- 文件操作：file_search, file_grep, file_diff, zip_list
- 系统信息：process_info, hash_generate, url_parse
- Web辅助：html_extract, markdown_render
- 图片信息：image_info
"""

import asyncio
import csv
import hashlib
import io
import json
import logging
import os
import re
import sqlite3
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import urlparse, parse_qs

from app.plugins.base import Plugin, Tool

logger = logging.getLogger(__name__)


class EnhancedToolsPlugin(Plugin):
    """增强工具集 — 文档、数据、文本、文件、系统工具"""

    name = "enhanced_tools"
    description = "增强工具集：文档解析、数据处理、文本分析、文件操作、系统信息"
    version = "1.0.0"
    author = "Fugue Team"
    tags = ["documents", "data", "text", "files", "system", "utility"]

    # ──────────────────── 文档解析工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "DOCX文件路径"},
                "max_paragraphs": {"type": "integer", "description": "最大段落数（默认全部）"}
            },
            "required": ["path"],
        },
        permissions="safe", category="document", version="1.0.0"
    )
    async def docx_read(self, path: str, max_paragraphs: int = None) -> str:
        """读取 Word (.docx) 文档内容，提取文本和表格"""
        def _read():
            try:
                from docx import Document
            except ImportError:
                return "[错误] python-docx 未安装，请运行: pip install python-docx"

            if not os.path.exists(path):
                return f"[错误] 文件不存在: {path}"

            try:
                doc = Document(path)
                parts = []

                # 提取段落
                paragraphs = doc.paragraphs
                if max_paragraphs:
                    paragraphs = paragraphs[:max_paragraphs]

                for para in paragraphs:
                    text = para.text.strip()
                    if text:
                        # 根据样式标记标题
                        if para.style.name.startswith('Heading'):
                            level = para.style.name.replace('Heading ', '')
                            parts.append(f"{'#' * int(level)} {text}")
                        else:
                            parts.append(text)

                # 提取表格
                for i, table in enumerate(doc.tables):
                    parts.append(f"\n[表格 {i+1}]")
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        parts.append(" | ".join(cells))

                content = "\n".join(parts)
                if not content.strip():
                    return "[提示] 文档内容为空"

                # 截断过长内容
                if len(content) > 100000:
                    content = content[:100000] + f"\n\n... [文档过长，已截断，共 {len(content)} 字符]"

                return f"[DOCX文档] {os.path.basename(path)}\n段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}\n\n{content}"

            except Exception as e:
                return f"[错误] 读取DOCX失败: {e}"

        return await asyncio.to_thread(_read)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Excel文件路径"},
                "sheet_name": {"type": "string", "description": "工作表名称（默认第一个）"},
                "max_rows": {"type": "integer", "description": "最大行数（默认100）"}
            },
            "required": ["path"],
        },
        permissions="safe", category="document", version="1.0.0"
    )
    async def xlsx_read(self, path: str, sheet_name: str = None, max_rows: int = 100) -> str:
        """读取 Excel (.xlsx) 文件，返回指定工作表的数据"""
        def _read():
            try:
                from openpyxl import load_workbook
            except ImportError:
                return "[错误] openpyxl 未安装，请运行: pip install openpyxl"

            if not os.path.exists(path):
                return f"[错误] 文件不存在: {path}"

            try:
                wb = load_workbook(path, read_only=True, data_only=True)
                sheets = wb.sheetnames

                if sheet_name:
                    if sheet_name not in sheets:
                        return f"[错误] 工作表 '{sheet_name}' 不存在。可用: {', '.join(sheets)}"
                    ws = wb[sheet_name]
                else:
                    ws = wb.active

                rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= max_rows:
                        break
                    cells = [str(c) if c is not None else "" for c in row]
                    rows.append(cells)

                wb.close()

                if not rows:
                    return f"[提示] 工作表为空"

                # 格式化为表格
                header = rows[0] if rows else []
                data = rows[1:] if len(rows) > 1 else []
                output = f"[Excel] {os.path.basename(path)} | 工作表: {ws.title}\n"
                output += f"行数: {ws.max_row}, 列数: {ws.max_column}\n\n"
                output += " | ".join(header) + "\n"
                output += "-" * 60 + "\n"
                for row in data:
                    output += " | ".join(row) + "\n"

                if ws.max_row and ws.max_row > max_rows:
                    output += f"\n... 共 {ws.max_row} 行，仅显示前 {max_rows} 行"

                return output
            except Exception as e:
                return f"[错误] 读取Excel失败: {e}"

        return await asyncio.to_thread(_read)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "CSV文件路径"},
                "encoding": {"type": "string", "description": "文件编码（默认自动检测）"},
                "max_rows": {"type": "integer", "description": "最大行数（默认100）"},
                "delimiter": {"type": "string", "description": "分隔符（默认自动检测）"}
            },
            "required": ["path"],
        },
        permissions="safe", category="document", version="1.0.0"
    )
    async def csv_analyze(self, path: str, encoding: str = None, max_rows: int = 100, delimiter: str = None) -> str:
        """读取并分析 CSV 文件，返回数据摘要和预览"""
        def _analyze():
            if not os.path.exists(path):
                return f"[错误] 文件不存在: {path}"

            try:
                # 自动检测编码
                raw = open(path, 'rb').read()
                if encoding:
                    text = raw.decode(encoding)
                else:
                    import chardet
                    detected = chardet.detect(raw)
                    enc = detected.get('encoding', 'utf-8')
                    text = raw.decode(enc, errors='replace')

                # 自动检测分隔符
                if delimiter:
                    sep = delimiter
                else:
                    sniffer = csv.Sniffer()
                    try:
                        dialect = sniffer.sniff(text[:2048])
                        sep = dialect.delimiter
                    except csv.Error:
                        sep = ','

                reader = csv.reader(io.StringIO(text), delimiter=sep)
                all_rows = list(reader)

                if not all_rows:
                    return "[提示] CSV文件为空"

                header = all_rows[0]
                data = all_rows[1:]
                total_rows = len(data)

                # 数据统计
                output = f"[CSV分析] {os.path.basename(path)}\n"
                output += f"分隔符: '{sep}' | 列数: {len(header)} | 数据行数: {total_rows}\n"
                output += f"列名: {', '.join(header)}\n\n"

                # 数值列统计
                numeric_cols = []
                for col_idx in range(len(header)):
                    values = []
                    for row in data[:max_rows]:
                        if col_idx < len(row):
                            try:
                                values.append(float(row[col_idx]))
                            except (ValueError, IndexError):
                                pass
                    if len(values) > len(data[:max_rows]) * 0.5:
                        numeric_cols.append(col_idx)
                        if values:
                            output += f"列 '{header[col_idx]}': min={min(values):.2f}, max={max(values):.2f}, avg={sum(values)/len(values):.2f}\n"

                output += "\n"

                # 数据预览
                preview_rows = data[:max_rows]
                output += " | ".join(header) + "\n"
                output += "-" * 60 + "\n"
                for row in preview_rows:
                    output += " | ".join(row[:len(header)]) + "\n"

                if total_rows > max_rows:
                    output += f"\n... 共 {total_rows} 行，仅显示前 {max_rows} 行"

                return output
            except Exception as e:
                return f"[错误] 读取CSV失败: {e}"

        return await asyncio.to_thread(_analyze)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "PDF文件路径"},
                "max_pages": {"type": "integer", "description": "最大页数（默认全部）"}
            },
            "required": ["path"],
        },
        permissions="safe", category="document", version="1.0.0"
    )
    async def pdf_read(self, path: str, max_pages: int = None) -> str:
        """读取 PDF 文件，提取文本内容"""
        def _read():
            try:
                import pdfplumber
            except ImportError:
                return "[错误] pdfplumber 未安装，请运行: pip install pdfplumber"

            if not os.path.exists(path):
                return f"[错误] 文件不存在: {path}"

            try:
                parts = []
                with pdfplumber.open(path) as pdf:
                    total_pages = len(pdf.pages)
                    pages = pdf.pages[:max_pages] if max_pages else pdf.pages

                    for i, page in enumerate(pages):
                        text = page.extract_text()
                        if text:
                            parts.append(f"--- 第 {i+1}/{total_pages} 页 ---\n{text}")

                        # 提取表格
                        tables = page.extract_tables()
                        for t_idx, table in enumerate(tables):
                            parts.append(f"[第{i+1}页 表格{t_idx+1}]")
                            for row in table:
                                cells = [str(c or "") for c in row]
                                parts.append(" | ".join(cells))

                content = "\n\n".join(parts)
                if not content.strip():
                    return f"[提示] PDF 未提取到文本（可能是扫描型PDF，需OCR）"

                if len(content) > 100000:
                    content = content[:100000] + f"\n\n... [内容过长，已截断]"

                return f"[PDF文档] {os.path.basename(path)}\n总页数: {total_pages}\n\n{content}"
            except Exception as e:
                return f"[错误] 读取PDF失败: {e}"

        return await asyncio.to_thread(_read)

    # ──────────────────── 文档创建工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "保存路径（如 C:/Users/HP/Desktop/output.docx）"},
                "title": {"type": "string", "description": "文档标题"},
                "paragraphs": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "段落内容列表"
                },
                "headings": {
                    "type": "array",
                    "items": {"type": "object", "properties": {"level": {"type": "integer"}, "text": {"type": "string"}}},
                    "description": "标题列表 [{level: 1, text: '...'}]"
                },
                "table_data": {
                    "type": "array",
                    "items": {"type": "array", "items": {"type": "string"}},
                    "description": "表格数据（二维数组，第一行为表头）"
                }
            },
            "required": ["path"],
        },
        permissions="caution", category="document", version="1.0.0"
    )
    async def docx_create(self, path: str, title: str = None, paragraphs: list = None, headings: list = None, table_data: list = None) -> str:
        """创建 Word (.docx) 文档"""
        # 解析路径：相对路径基于当前工作空间或用户桌面
        import os as _os
        if not _os.path.isabs(path):
            from app.engine.tools import _workspace_dir as _ws
            base = _ws or _os.path.expanduser("~/Desktop")
            path = _os.path.join(base, path)
        path = _os.path.abspath(path)

        def _create():
            try:
                from docx import Document
                from docx.shared import Pt, Inches
            except ImportError:
                return "[错误] python-docx 未安装"

            try:
                doc = Document()

                # 设置默认字体
                style = doc.styles['Normal']
                font = style.font
                font.name = '宋体'
                font.size = Pt(12)

                if title:
                    doc.add_heading(title, level=0)

                if headings:
                    for h in headings:
                        level = min(h.get('level', 1), 4)
                        doc.add_heading(h.get('text', ''), level=level)

                if paragraphs:
                    for p in paragraphs:
                        doc.add_paragraph(p)

                if table_data and len(table_data) > 0:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            if j < len(table.rows[i].cells):
                                table.rows[i].cells[j].text = str(cell_text)

                # 确保目录存在
                os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
                doc.save(path)

                return f"✅ Word文档创建成功: {path}\n段落数: {len(paragraphs or [])}, 表格: {'有' if table_data else '无'}"
            except Exception as e:
                return f"[错误] 创建DOCX失败: {e}"

        return await asyncio.to_thread(_create)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "format": {"type": "string", "description": "日期格式（默认 %Y-%m-%d %H:%M:%S）", "default": "%Y-%m-%d %H:%M:%S"},
                "timezone": {"type": "string", "description": "时区（如 Asia/Shanghai）"}
            },
        },
        permissions="safe", category="system", version="1.0.0"
    )
    async def date_time(self, format: str = "%Y-%m-%d %H:%M:%S", timezone: str = None) -> str:
        """获取当前日期和时间"""
        from datetime import datetime, timezone as tz
        if timezone:
            try:
                import pytz
                tz_obj = pytz.timezone(timezone)
                now = datetime.now(tz_obj)
            except ImportError:
                now = datetime.now()
        else:
            now = datetime.now()
        return f"[日期时间] {now.strftime(format)}"

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "保存路径"},
                "rows": {
                    "type": "array",
                    "items": {"type": "array", "items": {"type": "string"}},
                    "description": "数据行（二维数组，第一行为表头）"
                },
                "encoding": {"type": "string", "description": "编码（默认utf-8-sig，Excel兼容）", "default": "utf-8-sig"}
            },
            "required": ["path", "rows"],
        },
        permissions="caution", category="document", version="1.0.0"
    )
    async def csv_create(self, path: str, rows: list, encoding: str = "utf-8-sig") -> str:
        """创建 CSV 文件"""
        import os as _os
        if not _os.path.isabs(path):
            from app.engine.tools import _workspace_dir as _ws
            base = _ws or _os.path.expanduser("~/Desktop")
            path = _os.path.join(base, path)
        path = _os.path.abspath(path)

        def _create():
            try:
                import csv
                os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
                with open(path, 'w', newline='', encoding=encoding) as f:
                    writer = csv.writer(f)
                    for row in rows:
                        writer.writerow(row)
                return f"✅ CSV文件创建成功: {path} ({len(rows)} 行)"
            except Exception as e:
                return f"[错误] 创建CSV失败: {e}"

        return await asyncio.to_thread(_create)

    # ──────────────────── 数据处理工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "data": {"type": "string", "description": "JSON字符串"},
                "query": {"type": "string", "description": "查询表达式（支持 key.subkey 和 [0] 索引）"},
                "filter_key": {"type": "string", "description": "过滤键名"},
                "filter_value": {"type": "string", "description": "过滤值"},
                "max_results": {"type": "integer", "description": "最大结果数（默认50）"}
            },
            "required": ["data"],
        },
        permissions="safe", category="data", version="1.0.0"
    )
    async def json_query(self, data: str, query: str = None, filter_key: str = None, filter_value: str = None, max_results: int = 50) -> str:
        """查询和过滤JSON数据"""
        def _process():
            try:
                obj = json.loads(data)
            except json.JSONDecodeError as e:
                return f"[错误] JSON解析失败: {e}"

            result = obj

            # 路径查询
            if query:
                for part in query.split('.'):
                    if part == '':
                        continue
                    if isinstance(result, list):
                        try:
                            idx = int(part)
                            result = result[idx] if idx < len(result) else None
                        except ValueError:
                            return f"[错误] 数组索引无效: {part}"
                    elif isinstance(result, dict):
                        result = result.get(part)
                    else:
                        return f"[错误] 无法查询: 当前值不是对象或数组"

                    if result is None:
                        return f"[结果] 路径 '{query}' 不存在"

            # 过滤
            if filter_key and isinstance(result, list):
                filtered = []
                for item in result:
                    if isinstance(item, dict):
                        val = str(item.get(filter_key, ''))
                        if filter_value is None or filter_value in val:
                            filtered.append(item)
                result = filtered[:max_results]

            # 格式化输出
            output_str = json.dumps(result, ensure_ascii=False, indent=2)
            if len(output_str) > 10000:
                output_str = output_str[:10000] + "\n... [输出过长，已截断]"

            summary = ""
            if isinstance(result, list):
                summary = f"结果数量: {len(result)}"
            elif isinstance(result, dict):
                summary = f"对象键: {', '.join(result.keys())[:200]}"

            return f"[JSON查询] {summary}\n\n{output_str}"

        return await asyncio.to_thread(_process)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "YAML文件路径"}
            },
            "required": ["path"],
        },
        permissions="safe", category="data", version="1.0.0"
    )
    async def yaml_parse(self, path: str) -> str:
        """解析 YAML 文件"""
        def _parse():
            try:
                import yaml
            except ImportError:
                return "[错误] pyyaml 未安装，请运行: pip install pyyaml"

            if not os.path.exists(path):
                return f"[错误] 文件不存在: {path}"

            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = yaml.safe_load(f)
                output = json.dumps(data, ensure_ascii=False, indent=2)
                if len(output) > 10000:
                    output = output[:10000] + "\n... [内容过长]"
                return f"[YAML] {os.path.basename(path)}\n\n{output}"
            except Exception as e:
                return f"[错误] 解析YAML失败: {e}"

        return await asyncio.to_thread(_parse)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "data": {"type": "string", "description": "数据（CSV格式或JSON数组）"},
                "format": {"type": "string", "enum": ["grid", "pipe", "markdown", "html"], "description": "输出格式", "default": "markdown"}
            },
            "required": ["data"],
        },
        permissions="safe", category="data", version="1.0.0"
    )
    async def table_format(self, data: str, format: str = "markdown") -> str:
        """将数据格式化为美观的表格"""
        def _format():
            try:
                from tabulate import tabulate
            except ImportError:
                return "[错误] tabulate 未安装，请运行: pip install tabulate"

            try:
                # 尝试JSON解析
                parsed = json.loads(data)
                if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
                    headers = list(parsed[0].keys())
                    rows = [[item.get(h, '') for h in headers] for item in parsed]
                elif isinstance(parsed, list) and parsed and isinstance(parsed[0], list):
                    headers = parsed[0]
                    rows = parsed[1:]
                else:
                    return "[错误] 不支持的数据格式，需要 [{...}, {...}] 或 [[...], [...]]"
            except json.JSONDecodeError:
                # 尝试CSV解析
                reader = csv.reader(io.StringIO(data))
                rows = list(reader)
                if not rows:
                    return "[错误] 数据为空"
                headers = rows[0]
                rows = rows[1:]

            fmt_map = {"grid": "grid", "pipe": "pipe", "markdown": "pipe", "html": "html"}
            table = tabulate(rows, headers=headers, tablefmt=format)
            return f"[表格] {len(rows)} 行 x {len(headers)} 列\n\n{table}"

        return await asyncio.to_thread(_format)

    # ──────────────────── 文本处理工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "待处理文本"},
                "pattern": {"type": "string", "description": "正则表达式"},
                "flags": {"type": "string", "description": "正则标志: i(忽略大小写), m(多行), s(点匹配换行)", "default": ""},
                "max_results": {"type": "integer", "description": "最大结果数（默认50）"}
            },
            "required": ["text", "pattern"],
        },
        permissions="safe", category="text", version="1.0.0"
    )
    async def regex_extract(self, text: str, pattern: str, flags: str = "", max_results: int = 50) -> str:
        """使用正则表达式提取文本中的数据"""
        def _extract():
            re_flags = 0
            if 'i' in flags: re_flags |= re.IGNORECASE
            if 'm' in flags: re_flags |= re.MULTILINE
            if 's' in flags: re_flags |= re.DOTALL

            try:
                matches = re.findall(pattern, text, re_flags)
            except re.error as e:
                return f"[错误] 正则表达式无效: {e}"

            if not matches:
                return f"[正则提取] 模式 '{pattern}' 未匹配到任何内容"

            # 处理分组匹配
            if isinstance(matches[0], tuple):
                results = [", ".join(m) for m in matches[:max_results]]
            else:
                results = [str(m) for m in matches[:max_results]]

            output = f"[正则提取] 模式: {pattern}\n匹配数: {len(matches)}\n\n"
            for i, r in enumerate(results, 1):
                output += f"{i}. {r}\n"

            if len(matches) > max_results:
                output += f"\n... 共 {len(matches)} 个匹配，仅显示前 {max_results} 个"

            return output

        return await asyncio.to_thread(_extract)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "text_a": {"type": "string", "description": "文本A"},
                "text_b": {"type": "string", "description": "文本B"},
                "context_lines": {"type": "integer", "description": "差异上下文行数（默认3）", "default": 3}
            },
            "required": ["text_a", "text_b"],
        },
        permissions="safe", category="text", version="1.0.0"
    )
    async def text_diff(self, text_a: str, text_b: str, context_lines: int = 3) -> str:
        """比较两段文本的差异"""
        def _diff():
            import difflib

            lines_a = text_a.splitlines(keepends=True)
            lines_b = text_b.splitlines(keepends=True)

            diff = list(difflib.unified_diff(
                lines_a, lines_b,
                fromfile="文本A", tofile="文本B",
                n=context_lines
            ))

            if not diff:
                return "[文本差异] 两段文本完全相同"

            diff_text = "".join(diff)
            if len(diff_text) > 10000:
                diff_text = diff_text[:10000] + "\n... [差异过长，已截断]"

            return f"[文本差异]\n{diff_text}"

        return await asyncio.to_thread(_diff)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "待转换文本"},
                "operation": {
                    "type": "string",
                    "enum": ["upper", "lower", "title", "reverse", "trim", "strip_html", "word_count", "char_count", "base64_encode", "base64_decode", "url_encode", "url_decode"],
                    "description": "转换操作"
                }
            },
            "required": ["text", "operation"],
        },
        permissions="safe", category="text", version="1.0.0"
    )
    async def text_transform(self, text: str, operation: str) -> str:
        """对文本进行各种转换操作"""
        import base64
        from urllib.parse import quote, unquote
        from html.parser import HTMLParser

        class HTMLStripper(HTMLParser):
            def __init__(self):
                super().__init__()
                self.result = []
            def handle_data(self, d):
                self.result.append(d)
            def get_text(self):
                return ''.join(self.result)

        ops = {
            "upper": lambda t: t.upper(),
            "lower": lambda t: t.lower(),
            "title": lambda t: t.title(),
            "reverse": lambda t: t[::-1],
            "trim": lambda t: t.strip(),
            "strip_html": lambda t: (lambda p: (p.feed(t), p.get_text())[1])(HTMLStripper()),
            "word_count": lambda t: f"词数: {len(t.split())}, 行数: {len(t.splitlines())}",
            "char_count": lambda t: f"字符数: {len(t)}, 中文字符: {len(re.findall(r'[一-鿿]', t))}",
            "base64_encode": lambda t: base64.b64encode(t.encode()).decode(),
            "base64_decode": lambda t: base64.b64decode(t).decode(),
            "url_encode": lambda t: quote(t),
            "url_decode": lambda t: unquote(t),
        }

        func = ops.get(operation)
        if not func:
            return f"[错误] 不支持的操作: {operation}。可用: {', '.join(ops.keys())}"

        try:
            result = func(text)
            if len(str(result)) > 5000:
                result = str(result)[:5000] + "\n... [结果过长，已截断]"
            return f"[文本转换] 操作: {operation}\n\n{result}"
        except Exception as e:
            return f"[错误] 转换失败: {e}"

    # ──────────────────── 文件操作工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "directory": {"type": "string", "description": "搜索目录"},
                "pattern": {"type": "string", "description": "文件名模式（支持通配符 *.txt, *.docx 等）"},
                "recursive": {"type": "boolean", "description": "是否递归搜索子目录", "default": True},
                "max_results": {"type": "integer", "description": "最大结果数（默认50）"}
            },
            "required": ["directory", "pattern"],
        },
        permissions="safe", category="file", version="1.0.0"
    )
    async def file_search(self, directory: str, pattern: str, recursive: bool = True, max_results: int = 50) -> str:
        """在目录中搜索匹配模式的文件"""
        def _search():
            from fnmatch import fnmatch

            if not os.path.isdir(directory):
                return f"[错误] 目录不存在: {directory}"

            results = []
            if recursive:
                for root, dirs, files in os.walk(directory):
                    # 跳过隐藏目录和常见忽略目录
                    dirs[:] = [d for d in dirs if not d.startswith('.') and d not in ('node_modules', '__pycache__', '.git', 'venv')]
                    for f in files:
                        if fnmatch(f, pattern):
                            full = os.path.join(root, f)
                            size = os.path.getsize(full)
                            results.append({"name": f, "path": full, "size": size})
                            if len(results) >= max_results:
                                break
                    if len(results) >= max_results:
                        break
            else:
                for f in os.listdir(directory):
                    if fnmatch(f, pattern):
                        full = os.path.join(directory, f)
                        if os.path.isfile(full):
                            size = os.path.getsize(full)
                            results.append({"name": f, "path": full, "size": size})
                            if len(results) >= max_results:
                                break

            if not results:
                return f"[文件搜索] 目录 '{directory}' 中未找到匹配 '{pattern}' 的文件"

            output = f"[文件搜索] 目录: {directory}\n模式: {pattern}\n找到 {len(results)} 个文件\n\n"
            for r in results:
                size = r['size']
                if size < 1024: size_str = f"{size}B"
                elif size < 1024*1024: size_str = f"{size/1024:.1f}KB"
                else: size_str = f"{size/(1024*1024):.1f}MB"
                output += f"  {r['name']} ({size_str})\n    {r['path']}\n"

            return output

        return await asyncio.to_thread(_search)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "directory": {"type": "string", "description": "搜索目录"},
                "keyword": {"type": "string", "description": "搜索关键词（支持正则）"},
                "file_pattern": {"type": "string", "description": "限定文件类型（如 *.py, *.txt）", "default": "*"},
                "max_results": {"type": "integer", "description": "最大匹配数（默认30）"},
                "context_lines": {"type": "integer", "description": "上下文行数（默认0）", "default": 0}
            },
            "required": ["directory", "keyword"],
        },
        permissions="safe", category="file", version="1.0.0"
    )
    async def file_grep(self, directory: str, keyword: str, file_pattern: str = "*", max_results: int = 30, context_lines: int = 0) -> str:
        """在文件内容中搜索关键词（类似 grep）"""
        def _grep():
            from fnmatch import fnmatch

            if not os.path.isdir(directory):
                return f"[错误] 目录不存在: {directory}"

            results = []
            try:
                pattern = re.compile(keyword, re.IGNORECASE)
            except re.error:
                pattern = re.compile(re.escape(keyword), re.IGNORECASE)

            for root, dirs, files in os.walk(directory):
                dirs[:] = [d for d in dirs if not d.startswith('.') and d not in ('node_modules', '__pycache__', '.git', 'venv')]
                for fname in files:
                    if not fnmatch(fname, file_pattern):
                        continue
                    fpath = os.path.join(root, fname)
                    try:
                        with open(fpath, 'r', encoding='utf-8', errors='ignore') as f:
                            for line_num, line in enumerate(f, 1):
                                if pattern.search(line):
                                    results.append({
                                        "file": fpath,
                                        "line": line_num,
                                        "content": line.rstrip(),
                                    })
                                    if len(results) >= max_results:
                                        break
                    except (PermissionError, OSError):
                        continue
                    if len(results) >= max_results:
                        break

            if not results:
                return f"[Grep] 在 '{directory}' 中未找到 '{keyword}' 的匹配"

            output = f"[Grep] 目录: {directory}\n关键词: {keyword}\n匹配数: {len(results)}\n\n"
            for r in results:
                short_path = os.path.relpath(r['file'], directory)
                output += f"{short_path}:{r['line']}: {r['content']}\n"

            return output

        return await asyncio.to_thread(_grep)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "ZIP文件路径"},
                "list_only": {"type": "boolean", "description": "仅列出内容不解压", "default": True}
            },
            "required": ["path"],
        },
        permissions="safe", category="file", version="1.0.0"
    )
    async def zip_list(self, path: str, list_only: bool = True) -> str:
        """列出 ZIP 文件内容或解压"""
        def _zip():
            if not os.path.exists(path):
                return f"[错误] 文件不存在: {path}"

            try:
                with zipfile.ZipFile(path, 'r') as zf:
                    entries = zf.infolist()
                    if not entries:
                        return "[提示] ZIP文件为空"

                    output = f"[ZIP] {os.path.basename(path)}\n文件数: {len(entries)}\n\n"
                    total_size = 0
                    for entry in entries:
                        size = entry.file_size
                        total_size += size
                        if size < 1024: size_str = f"{size}B"
                        elif size < 1024*1024: size_str = f"{size/1024:.1f}KB"
                        else: size_str = f"{size/(1024*1024):.1f}MB"
                        prefix = "  [DIR] " if entry.is_dir() else "  "
                        output += f"{prefix}{entry.filename} ({size_str})\n"

                    if total_size < 1024*1024: total_str = f"{total_size/1024:.1f}KB"
                    else: total_str = f"{total_size/(1024*1024):.1f}MB"
                    output += f"\n总解压大小: {total_str}"

                    return output
            except zipfile.BadZipFile:
                return f"[错误] 无效的ZIP文件: {path}"
            except Exception as e:
                return f"[错误] 读取ZIP失败: {e}"

        return await asyncio.to_thread(_zip)

    # ──────────────────── 系统工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "要计算哈希的文本"},
                "algorithm": {"type": "string", "enum": ["md5", "sha1", "sha256", "sha512"], "description": "哈希算法", "default": "sha256"}
            },
            "required": ["text"],
        },
        permissions="safe", category="system", version="1.0.0"
    )
    async def hash_generate(self, text: str, algorithm: str = "sha256") -> str:
        """生成文本的哈希值"""
        data = text.encode('utf-8')
        h = hashlib.new(algorithm)
        h.update(data)
        return f"[哈希] 算法: {algorithm}\n输入长度: {len(text)} 字符\n哈希值: {h.hexdigest()}"

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "要解析的URL"}
            },
            "required": ["url"],
        },
        permissions="safe", category="system", version="1.0.0"
    )
    async def url_parse(self, url: str) -> str:
        """解析URL的各个组成部分"""
        parsed = urlparse(url)
        params = parse_qs(parsed.query)

        parts = [
            f"协议: {parsed.scheme}",
            f"主机: {parsed.hostname or 'N/A'}",
            f"端口: {parsed.port or '默认'}",
            f"路径: {parsed.path or '/'}",
            f"查询参数 ({len(params)} 个):",
        ]
        for k, v in params.items():
            parts.append(f"  {k} = {', '.join(v)}")
        if parsed.fragment:
            parts.append(f"锚点: {parsed.fragment}")

        return "[URL解析]\n" + "\n".join(parts)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {},
        },
        permissions="safe", category="system", version="1.0.0"
    )
    async def process_info(self) -> str:
        """获取当前系统和进程信息"""
        def _info():
            import platform
            info = [
                f"系统: {platform.system()} {platform.release()}",
                f"版本: {platform.version()}",
                f"架构: {platform.machine()}",
                f"Python: {platform.python_version()}",
                f"工作目录: {os.getcwd()}",
                f"CPU核心数: {os.cpu_count()}",
            ]

            # 磁盘空间
            try:
                import shutil
                total, used, free = shutil.disk_usage("/")
                info.append(f"磁盘空间: 总计 {total//(1024**3)}GB, 已用 {used//(1024**3)}GB, 可用 {free//(1024**3)}GB")
            except Exception:
                pass

            return "[系统信息]\n" + "\n".join(info)

        return await asyncio.to_thread(_info)

    # ──────────────────── Web辅助工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "html": {"type": "string", "description": "HTML内容"},
                "selector": {"type": "string", "description": "CSS选择器（可选，默认提取所有文本）"}
            },
            "required": ["html"],
        },
        permissions="safe", category="web", version="1.0.0"
    )
    async def html_extract(self, html: str, selector: str = None) -> str:
        """从HTML中提取文本内容"""
        def _extract():
            try:
                from bs4 import BeautifulSoup
            except ImportError:
                return "[错误] beautifulsoup4 未安装，请运行: pip install beautifulsoup4"

            soup = BeautifulSoup(html, 'html.parser')

            # 移除script和style
            for tag in soup(['script', 'style', 'noscript']):
                tag.decompose()

            if selector:
                elements = soup.select(selector)
                if not elements:
                    return f"[HTML提取] CSS选择器 '{selector}' 未匹配到任何元素"
                text = "\n".join(el.get_text(separator='\n', strip=True) for el in elements)
            else:
                text = soup.get_text(separator='\n', strip=True)

            if len(text) > 10000:
                text = text[:10000] + "\n... [内容过长，已截断]"

            return f"[HTML提取] 提取文本长度: {len(text)} 字符\n\n{text}"

        return await asyncio.to_thread(_extract)

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "markdown_text": {"type": "string", "description": "Markdown文本"}
            },
            "required": ["markdown_text"],
        },
        permissions="safe", category="web", version="1.0.0"
    )
    async def markdown_render(self, markdown_text: str) -> str:
        """将Markdown转换为HTML"""
        def _render():
            try:
                import markdown
            except ImportError:
                return "[错误] markdown 未安装，请运行: pip install markdown"

            html = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code', 'toc'])
            return f"[Markdown→HTML] 输出长度: {len(html)} 字符\n\n{html}"

        return await asyncio.to_thread(_render)

    # ──────────────────── 图片工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "图片文件路径"}
            },
            "required": ["path"],
        },
        permissions="safe", category="media", version="1.0.0"
    )
    async def image_info(self, path: str) -> str:
        """获取图片文件的元数据信息"""
        def _info():
            try:
                from PIL import Image
                from PIL.ExifTags import TAGS
            except ImportError:
                return "[错误] Pillow 未安装，请运行: pip install Pillow"

            if not os.path.exists(path):
                return f"[错误] 文件不存在: {path}"

            try:
                with Image.open(path) as img:
                    info = [
                        f"格式: {img.format}",
                        f"尺寸: {img.width} x {img.height}",
                        f"色彩模式: {img.mode}",
                        f"文件大小: {os.path.getsize(path) / 1024:.1f} KB",
                    ]

                    # EXIF 信息
                    exif = img.getexif()
                    if exif:
                        info.append("\nEXIF信息:")
                        for tag_id, value in list(exif.items())[:20]:
                            tag = TAGS.get(tag_id, tag_id)
                            info.append(f"  {tag}: {value}")

                    return "[图片信息] " + os.path.basename(path) + "\n" + "\n".join(info)
            except Exception as e:
                return f"[错误] 读取图片信息失败: {e}"

        return await asyncio.to_thread(_info)

    # ──────────────────── 数据库查询工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "db_path": {"type": "string", "description": "SQLite数据库文件路径"},
                "query": {"type": "string", "description": "SQL查询（仅允许SELECT）"},
                "max_rows": {"type": "integer", "description": "最大返回行数（默认50）"}
            },
            "required": ["db_path", "query"],
        },
        permissions="restricted", category="data", version="1.0.0"
    )
    async def sqlite_query(self, db_path: str, query: str, max_rows: int = 50) -> str:
        """查询本地SQLite数据库（仅允许SELECT）"""
        def _query():
            if not os.path.exists(db_path):
                return f"[错误] 数据库不存在: {db_path}"

            q = query.strip().upper()
            if not q.startswith("SELECT"):
                return "[错误] 安全限制：仅允许 SELECT 查询"
            if ";" in query.rstrip(";").rstrip():
                return "[错误] 不允许多条SQL语句"

            try:
                conn = sqlite3.connect(db_path)
                cursor = conn.execute(query)
                columns = [desc[0] for desc in cursor.description] if cursor.description else []
                rows = cursor.fetchmany(max_rows)
                conn.close()

                if not rows:
                    return f"[SQLite查询] 无结果\nSQL: {query}"

                output = f"[SQLite查询] 返回 {len(rows)} 行\n列: {', '.join(columns)}\n\n"
                output += " | ".join(columns) + "\n"
                output += "-" * 60 + "\n"
                for row in rows:
                    output += " | ".join(str(v) if v is not None else "NULL" for v in row) + "\n"

                return output
            except Exception as e:
                return f"[错误] 查询失败: {e}"

        return await asyncio.to_thread(_query)

    # ──────────────────── 文本统计工具 ────────────────────

    @Tool(
        input_schema={
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "待分析文本"},
                "top_n": {"type": "integer", "description": "返回前N个高频词（默认20）"}
            },
            "required": ["text"],
        },
        permissions="safe", category="text", version="1.0.0"
    )
    async def word_frequency(self, text: str, top_n: int = 20) -> str:
        """分析文本的词频统计"""
        def _analyze():
            # 中文分词（简易按字/词）
            if re.search(r'[一-鿿]', text):
                # 中文：按2-4字组合分词
                cn_words = re.findall(r'[一-鿿]{2,6}', text)
                stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这', '他', '她', '它', '们', '那', '些', '什么', '怎么', '如何', '可以', '但是', '因为', '所以', '如果', '虽然'}
                words = [w for w in cn_words if w not in stop_words]
            else:
                words = re.findall(r'[a-zA-Z]{3,}', text.lower())
                stop_words = {'the', 'is', 'at', 'which', 'on', 'and', 'a', 'an', 'to', 'in', 'for', 'of', 'with', 'that', 'this', 'it', 'from', 'by', 'as', 'are', 'was', 'were', 'be', 'been', 'has', 'have', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can'}
                words = [w for w in words if w not in stop_words]

            if not words:
                return "[词频分析] 未提取到有效词汇"

            freq = Counter(words).most_common(top_n)

            total_chars = len(text)
            total_words = len(words)

            output = f"[词频分析]\n总字符: {total_chars}, 有效词: {total_words}\n\n排名  词语        出现次数  占比\n"
            output += "-" * 40 + "\n"
            for i, (word, count) in enumerate(freq, 1):
                pct = count / total_words * 100
                output += f"{i:>3}.  {word:<10}  {count:>5}    {pct:.1f}%\n"

            return output

        return await asyncio.to_thread(_analyze)

    async def setup(self):
        logger.info("EnhancedToolsPlugin v%s initialized with %d tools", self.version, len(self.tools))

    async def cleanup(self):
        logger.info("EnhancedToolsPlugin cleanup")

    async def health_check(self) -> Dict[str, Any]:
        return {"healthy": True, "message": f"EnhancedToolsPlugin v{self.version} running"}


__all__ = ["EnhancedToolsPlugin"]
