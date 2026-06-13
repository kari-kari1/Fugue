"""执行引擎 — 支持并行/顺序执行、重试、超时"""

import asyncio
import json
import logging
import os
import re
import time
import traceback
import uuid
from collections import defaultdict
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, Any, List, Optional, Set

from simpleeval import simple_eval, InvalidExpression

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from sqlalchemy.orm.attributes import flag_modified

from app.core.database import db_session_manager
from app.models.crew import Crew, ProcessType
from app.models.agent import Agent
from app.models.task import Task
from app.models.condition import ConditionBranch
from app.models.loop import LoopConfig
from app.models.execution import (
    Execution, ExecutionStatus,
    TaskExecution, TaskExecutionStatus,
)
from app.models.human_review import HumanReviewConfig, HumanReviewRequest
from app.engine.llm_provider import get_llm_provider, LLMResponse, ToolCall
from app.engine.tools import get_openai_tools, get_anthropic_tools, execute_tool, get_plugin_tool_schemas
from app.engine.checkpoint_manager import CheckpointManager
from app.services.event_publisher import event_publisher

logger = logging.getLogger(__name__)


# 辅助函数：获取数据库会话
def get_db_session():
    """获取适用于当前事件循环的数据库会话"""
    return db_session_manager.get_session()


class ExecutionEngine:
    """执行引擎"""

    def __init__(self, execution_id: str, llm_api_keys: Dict[str, str] = None, llm_base_urls: Dict[str, str] = None, max_turns: int = 10):
        self.execution_id = execution_id
        self.llm_api_keys = llm_api_keys or {}
        self.llm_base_urls = llm_base_urls or {}
        self.max_turns = max_turns
        self._cancelled = False
        self._paused = False
        self._iteration_model = None  # 迭代优化时使用的模型名称

    def cancel(self):
        self._cancelled = True

    def pause(self):
        self._paused = True

    async def _wait_for_human_review(
        self,
        review_config: 'HumanReviewConfig',
        execution_id: str,
        task_id: str,
        context: Dict[str, Any],
    ) -> Any:
        """等待人工审核完成

        Args:
            review_config: 审核配置
            execution_id: 当前执行ID
            task_id: 关联的任务/节点ID
            context: 执行上下文（用于格式化 prompt 模板）

        Returns:
            审核通过时返回 review_result（dict 或 None）
        """
        # 创建审核请求
        async with get_db_session() as db:
            review_request = HumanReviewRequest(
                execution_id=execution_id,
                task_id=task_id,
                review_type=review_config.review_type,
                prompt=self._format_prompt(review_config.prompt, context),
                options=review_config.options,
                timeout_action=review_config.timeout_action,
            )

            # 设置超时时间（默认24小时）
            timeout_seconds = review_config.timeout_seconds or 86400
            review_request.timeout_at = datetime.utcnow() + timedelta(
                seconds=timeout_seconds
            )

            db.add(review_request)

            # 更新执行状态为等待审核
            execution = await db.get(Execution, execution_id)
            if execution:
                execution.status = ExecutionStatus.WAITING_REVIEW
                self._add_trace(
                    execution, "review.requested",
                    data={
                        "review_id": None,  # commit 后填充
                        "review_type": review_config.review_type,
                        "prompt": review_request.prompt[:200],
                    },
                )

            await db.commit()
            await db.refresh(review_request)

            # 更新 trace 中的 review_id
            if execution and execution.trace:
                for event in reversed(execution.trace):
                    if event.get("event_type") == "review.requested" and event["data"].get("review_id") is None:
                        event["data"]["review_id"] = review_request.id
                        flag_modified(execution, "trace")
                        await db.commit()
                        break

        review_request_id = review_request.id
        logger.info(
            f"[EXECUTOR] Waiting for human review {review_request_id} "
            f"(execution={execution_id}, type={review_config.review_type})"
        )

        # 发布审核等待事件
        await event_publisher.publish_progress(
            execution_id=execution_id,
            completed=0,
            total=0,
        )

        # I3: 轮询等待审核结果（指数退避 + 超时上限）
        max_wait_seconds = timeout_seconds
        poll_start_time = datetime.utcnow()
        poll_interval = 2.0  # 起始间隔 2 秒
        poll_max_interval = 60.0  # 最大间隔 60 秒

        while True:
            await asyncio.sleep(poll_interval)
            poll_interval = min(poll_interval * 1.5, poll_max_interval)  # 指数退避

            if self._cancelled:
                raise Exception("执行已取消")

            # 检查总等待时间是否超过上限
            elapsed = (datetime.utcnow() - poll_start_time).total_seconds()
            if elapsed > max_wait_seconds:
                raise Exception(
                    f"人工审核等待超时（已等待 {int(elapsed)} 秒，上限 {max_wait_seconds} 秒）"
                )

            async with get_db_session() as db:
                request = await db.get(HumanReviewRequest, review_request_id)
                if not request:
                    raise Exception("审核请求不存在")

                if request.status == "approved":
                    logger.info(f"[EXECUTOR] Review {review_request_id} approved")
                    # 恢复执行状态
                    execution = await db.get(Execution, execution_id)
                    if execution:
                        execution.status = ExecutionStatus.RUNNING
                        self._add_trace(
                            execution, "review.approved",
                            data={"review_id": review_request_id},
                        )
                        await db.commit()
                    return request.review_result

                elif request.status == "rejected":
                    logger.info(f"[EXECUTOR] Review {review_request_id} rejected")
                    execution = await db.get(Execution, execution_id)
                    if execution:
                        execution.status = ExecutionStatus.FAILED
                        execution.error_log = f"人工审核被拒绝: {request.review_comment or ''}"
                        self._add_trace(
                            execution, "review.rejected",
                            data={
                                "review_id": review_request_id,
                                "comment": request.review_comment,
                            },
                        )
                        await db.commit()
                    raise Exception(f"人工审核被拒绝: {request.review_comment or ''}")

                elif request.status == "skipped":
                    logger.info(f"[EXECUTOR] Review {review_request_id} skipped")
                    execution = await db.get(Execution, execution_id)
                    if execution:
                        execution.status = ExecutionStatus.RUNNING
                        self._add_trace(
                            execution, "review.skipped",
                            data={"review_id": review_request_id},
                        )
                        await db.commit()
                    return None

                # 检查是否超时
                if request.timeout_at and datetime.utcnow() > request.timeout_at:
                    timeout_action = request.timeout_action or "reject"
                    logger.info(
                        f"[EXECUTOR] Review {review_request_id} timed out, "
                        f"action={timeout_action}"
                    )
                    request.status = "skipped"  # 统一标记为已处理
                    request.reviewed_at = datetime.utcnow()
                    execution = await db.get(Execution, execution_id)
                    if execution:
                        self._add_trace(
                            execution, "review.timeout",
                            data={
                                "review_id": review_request_id,
                                "action": timeout_action,
                            },
                        )
                    await db.commit()

                    if timeout_action == "approve":
                        if execution:
                            execution.status = ExecutionStatus.RUNNING
                            await db.commit()
                        return None
                    elif timeout_action == "skip":
                        if execution:
                            execution.status = ExecutionStatus.RUNNING
                            await db.commit()
                        return None
                    else:
                        if execution:
                            execution.status = ExecutionStatus.FAILED
                            execution.error_log = "人工审核超时"
                            await db.commit()
                        raise Exception("人工审核超时")

    @staticmethod
    async def submit_review(
        review_request_id: str,
        user_id: str,
        action: str,
        result: Any = None,
        comment: str = None,
    ):
        """提交审核结果（供外部API调用）

        Args:
            review_request_id: 审核请求ID
            user_id: 审核人ID
            action: 审核动作 ("approve", "reject", "skip")
            result: 审核结果（approval类型可传 None，input/selection 类型传具体内容）
            comment: 审核备注
        """
        async with get_db_session() as db:
            review_request = await db.get(HumanReviewRequest, review_request_id)
            if not review_request:
                raise ValueError("审核请求不存在")

            if review_request.status != "pending":
                raise ValueError(f"审核请求已处理（当前状态: {review_request.status}）")

            # 检查关联执行是否仍处于等待审核状态
            execution = await db.get(Execution, review_request.execution_id)
            if not execution or execution.status != ExecutionStatus.WAITING_REVIEW:
                raise ValueError("执行未处于等待审核状态")

            # action → status 映射（"approve" → "approved" 等）
            status_map = {"approve": "approved", "reject": "rejected", "skip": "skipped"}
            review_request.status = status_map.get(action, action)
            review_request.reviewer_id = user_id
            review_request.review_result = result
            review_request.review_comment = comment
            review_request.reviewed_at = datetime.utcnow()

            await db.commit()

            logger.info(
                f"[EXECUTOR] Review {review_request_id} submitted: "
                f"status={review_request.status}, reviewer={user_id}"
            )

    def _format_prompt(self, template: str, context: Dict[str, Any]) -> str:
        """格式化审核提示模板

        支持 {variable} 占位符，从 context 中取值。
        """
        try:
            return template.format(**context)
        except (KeyError, IndexError, ValueError):
            return template

    async def run(self, resume: bool = False):
        """执行工作流（支持断点续传）— I2: 包含 30 分钟总超时

        Args:
            resume: 是否从断点恢复执行
        """
        execution_id = self.execution_id
        logger.info(f"[EXECUTOR] Starting execution {execution_id}, resume={resume}")

        # I2: 30 分钟总超时保护
        try:
            return await asyncio.wait_for(
                self._run_inner(resume),
                timeout=1800,  # 30 分钟
            )
        except asyncio.TimeoutError:
            logger.error(f"[EXECUTOR] Execution {execution_id} timed out after 30 minutes")
            # 清理 worktree 和 workspace（超时路径）
            if getattr(self, '_worktree_path', None):
                try:
                    from app.services.worktree_manager import get_worktree_manager
                    wt_mgr = get_worktree_manager()
                    await wt_mgr.remove_worktree(self._worktree_path)
                except Exception as e:
                    logger.debug(f"Worktree cleanup failed (non-critical): {e}")
                self._worktree_path = None
            try:
                from app.engine.tools import clear_workspace_dir
                clear_workspace_dir()
            except Exception as e:
                logger.debug(f"Workspace cleanup failed (non-critical): {e}")
            try:
                async with get_db_session() as db:
                    execution = await db.get(Execution, execution_id)
                    if execution and execution.status == ExecutionStatus.RUNNING:
                        execution.status = ExecutionStatus.FAILED
                        execution.error_log = "执行超时（30分钟上限）"
                        await db.commit()
            except Exception as e:
                logger.error(f"Failed to mark timed-out execution: {e}")

    async def _run_inner(self, resume: bool = False):
        """实际执行逻辑（被 run() 包装以添加超时）"""
        execution_id = self.execution_id
        try:
            async with get_db_session() as db:
                logger.info(f"[EXECUTOR] Database session created for {execution_id}")
                execution = await db.get(Execution, execution_id)
                if not execution:
                    logger.error(f"[EXECUTOR] Execution {execution_id} not found in database")
                    return

                result = await db.execute(
                    select(Crew)
                    .where(Crew.id == execution.crew_id)
                    .options(
                        selectinload(Crew.agents),
                        selectinload(Crew.tasks),
                        selectinload(Crew.condition_branches),
                        selectinload(Crew.loop_configs),
                        selectinload(Crew.review_configs),
                        selectinload(Crew.memory_configs),
                    )
                )
                crew = result.scalar_one_or_none()
                if not crew:
                    await self._fail_execution(db, execution, "工作流不存在")
                    return

                # 提取记忆配置（供任务执行时使用）
                self._memory_config = crew.memory_configs[0] if crew.memory_configs else None
                self._cost_budget = crew.cost_budget  # 美元，None 表示不限
                self._workspace_dir = crew.workspace_dir  # 工作空间目录
                # 同步工作空间到工具模块
                if crew.workspace_dir:
                    from app.engine.tools import set_workspace_dir
                    set_workspace_dir(crew.workspace_dir)
                    logger.info(f"[EXECUTOR] Workspace set to: {crew.workspace_dir}")

                # 初始化断点管理器
                checkpoint_manager = CheckpointManager(db, execution_id)

                # 恢复状态（如果 resume=True）
                resumed_outputs: Dict[str, str] = {}
                resumed_completed_ids: List[str] = []
                resumed_total_tokens = 0
                resumed_total_cost = 0.0

                if resume:
                    checkpoint_data = await checkpoint_manager.resume_from_checkpoint()
                    if checkpoint_data:
                        resumed_completed_ids = checkpoint_data["completed_task_ids"]
                        resumed_outputs = checkpoint_data["task_outputs"]
                        resumed_total_tokens = checkpoint_data["total_tokens"]
                        resumed_total_cost = checkpoint_data["total_cost_usd"]
                        logger.info(
                            f"[EXECUTOR] Resumed from checkpoint: "
                            f"{len(resumed_completed_ids)} tasks already completed"
                        )
                    else:
                        logger.warning(
                            f"[EXECUTOR] resume=True but no checkpoint found, starting fresh"
                        )
                        resume = False

                # 标记开始
                execution.status = ExecutionStatus.RUNNING
                execution.started_at = datetime.utcnow()
                execution.trace = [] if not resume else (execution.trace or [])
                trace_event = "crew.resumed" if resume else "crew.started"
                self._add_trace(
                    execution, trace_event,
                    data={"crew_name": crew.name, "process": crew.process.value},
                )
                await db.commit()

                # 检测沙箱类型并记录
                try:
                    from app.engine.sandbox import get_sandbox_manager
                    _sandbox = get_sandbox_manager()
                    _sandbox_type = _sandbox._auto_select_sandbox().value if hasattr(_sandbox, '_auto_select_sandbox') else 'none'
                    execution.sandbox_type = _sandbox_type
                    await db.commit()
                except Exception:
                    execution.sandbox_type = 'none'

                # Git Worktree 隔离执行
                if crew.workspace_dir:
                    try:
                        from app.services.worktree_manager import get_worktree_manager
                        import os as _os
                        wt_mgr = get_worktree_manager()
                        # 检测是否为 git 仓库，非 git 时跳过 worktree 创建
                        repo_path = crew.workspace_dir
                        git_dir = _os.path.join(repo_path, ".git")
                        if _os.path.exists(git_dir):
                            self._worktree_path = await wt_mgr.create_worktree(
                                repo_path=repo_path,
                                worktree_name=f"exec-{execution_id[:8]}",
                            )
                        else:
                            # 非 git 仓库：直接使用 workspace_dir 作为隔离目录
                            self._worktree_path = _os.path.join(
                                repo_path, ".fugue", f"exec-{execution_id[:8]}"
                            )
                            _os.makedirs(self._worktree_path, exist_ok=True)
                        execution.worktree_path = self._worktree_path
                        await db.commit()
                        from app.engine.tools import set_workspace_dir
                        set_workspace_dir(self._worktree_path)
                        logger.info(f"[EXECUTOR] Worktree/Dir created: {self._worktree_path}")
                    except Exception as wt_err:
                        logger.warning(f"[EXECUTOR] Worktree creation failed, using original workspace: {wt_err}")

                # 发布工作流开始事件
                await event_publisher.publish_crew_started(
                    execution_id=execution_id,
                    crew_name=crew.name,
                    process_type=crew.process.value,
                )

                agents_map = {a.id: a for a in (crew.agents or [])}
                all_tasks = crew.tasks or []
                tasks_with_agent = [t for t in all_tasks if t.agent_id]
                task_map = {t.id: t for t in tasks_with_agent}

                # 条件分支过滤：根据条件表达式决定要执行的任务子集
                if crew.condition_branches:
                    condition_result = await self.execute_with_conditions(
                        tasks=tasks_with_agent,
                        conditions=list(crew.condition_branches),
                        context=resumed_outputs if resume else {},
                    )
                    executable_tasks = condition_result["executable_tasks"]
                    skipped_tasks = condition_result["skipped_tasks"]
                    tasks_with_agent = executable_tasks
                    task_map = {t.id: t for t in tasks_with_agent}

                    # 记录被条件分支跳过的任务
                    for skipped in skipped_tasks:
                        self._add_trace(execution, "task.skipped", task_name=skipped.name,
                                        data={"reason": "条件分支跳过"})

                    self._add_trace(execution, "crew.condition_evaluated", data={
                        "total_conditions": len(crew.condition_branches),
                        "executable_tasks": len(executable_tasks),
                        "skipped_tasks": len(skipped_tasks),
                    })

                    logger.info(
                        f"[EXECUTOR] Condition branches evaluated: "
                        f"{len(executable_tasks)} executable, {len(skipped_tasks)} skipped"
                    )

                # 创建或恢复TaskExecution记录（需在循环执行之前）
                task_exec_map: Dict[str, TaskExecution] = {}
                if resume:
                    # 恢复：查询已有的 TaskExecution 记录
                    existing_te_result = await db.execute(
                        select(TaskExecution).where(
                            TaskExecution.execution_id == execution.id
                        )
                    )
                    existing_te_map = {
                        te.task_id: te for te in existing_te_result.scalars().all()
                    }
                    for task in tasks_with_agent:
                        if task.id in existing_te_map:
                            te = existing_te_map[task.id]
                            if task.id in set(resumed_completed_ids) and te.status != TaskExecutionStatus.COMPLETED:
                                te.status = TaskExecutionStatus.COMPLETED
                            task_exec_map[task.id] = te
                        else:
                            te = TaskExecution(
                                execution_id=execution.id,
                                task_id=task.id,
                                agent_id=task.agent_id,
                                status=TaskExecutionStatus.PENDING,
                            )
                            db.add(te)
                            task_exec_map[task.id] = te
                    await db.commit()
                else:
                    # 新建：为所有任务创建 TaskExecution 记录
                    for task in tasks_with_agent:
                        te = TaskExecution(
                            execution_id=execution.id,
                            task_id=task.id,
                            agent_id=task.agent_id,
                            status=TaskExecutionStatus.PENDING,
                        )
                        db.add(te)
                        task_exec_map[task.id] = te
                    await db.commit()

                # 初始化或恢复执行状态
                task_outputs: Dict[str, str] = dict(resumed_outputs) if resume else {}
                total_tokens = resumed_total_tokens if resume else 0
                total_cost = resumed_total_cost if resume else 0.0
                total_tasks = len(tasks_with_agent)
                completed_count = len(resumed_completed_ids) if resume else 0
                completed_task_ids: Set[str] = set(resumed_completed_ids) if resume else set()

                # 循环执行：如有循环配置，先执行循环体任务
                loop_task_ids: Set[str] = set()
                loop_outputs: Dict[str, Any] = {}
                if crew.loop_configs:
                    loop_configs = list(crew.loop_configs)
                    for lc in loop_configs:
                        loop_task_ids.update(lc.loop_body_task_ids or [])

                    # 分离循环任务和非循环任务
                    non_loop_tasks = [t for t in tasks_with_agent if t.id not in loop_task_ids]
                    loop_tasks_list = [t for t in tasks_with_agent if t.id in loop_task_ids]

                    if loop_tasks_list:
                        self._add_trace(execution, "crew.loop_start", data={
                            "loop_configs": len(loop_configs),
                            "loop_tasks": len(loop_tasks_list),
                        })

                        # 执行循环
                        loop_result = await self.execute_with_loops(
                            tasks=tasks_with_agent,
                            loop_configs=loop_configs,
                            agents_map=agents_map,
                            execution=execution,
                            task_outputs=task_outputs,
                            task_exec_map=task_exec_map,
                            db=db,
                            checkpoint_manager=checkpoint_manager,
                            total_tokens=total_tokens,
                            total_cost=total_cost,
                            completed_task_ids=completed_task_ids,
                            completed_count=completed_count,
                            total_tasks=total_tasks,
                        )

                        loop_outputs = loop_result.get("loop_results", [])
                        total_tokens = loop_result["total_tokens"]
                        total_cost = loop_result["total_cost"]
                        completed_count = loop_result["completed_count"]

                        # 将循环结果写入上下文
                        task_outputs["__loop_results__"] = str(loop_outputs)

                        self._add_trace(execution, "crew.loop_end", data={
                            "loop_results_count": len(loop_outputs),
                        })

                    # 后续只处理非循环任务
                    tasks_with_agent = non_loop_tasks
                    task_map = {t.id: t for t in tasks_with_agent}

                is_parallel = crew.process == ProcessType.PARALLEL
                is_orchestrator = crew.process == ProcessType.ORCHESTRATOR
                is_event_flow = crew.process == ProcessType.EVENT_FLOW
                is_plan_execute = crew.process == ProcessType.PLAN_EXECUTE
                is_hierarchical = crew.process == ProcessType.HIERARCHICAL
                is_evaluator_optimizer = crew.process == ProcessType.EVALUATOR_OPTIMIZER
                is_prompt_chain = crew.process == ProcessType.PROMPT_CHAIN
                is_router = crew.process == ProcessType.ROUTER

                # Plan-Execute 模式：Agent 自主生成计划并逐步执行
                if is_plan_execute:
                    pe_result = await self._execute_plan_execute(
                        db, execution, tasks_with_agent, agents_map, task_map,
                        task_outputs, task_exec_map, completed_task_ids,
                        total_tokens, total_cost, completed_count, total_tasks,
                    )
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.total_tokens_used = pe_result["total_tokens"]
                    execution.total_cost_usd = pe_result["total_cost"]
                    execution.results = {"outputs": task_outputs, "mode": "plan_execute", **pe_result}
                    self._add_trace(execution, "crew.completed", data={"mode": "plan_execute"})
                    await db.commit()
                    return

                # Hierarchical 模式：Manager 审查 Worker 输出
                if is_hierarchical and len(agents_map) >= 2 and tasks_with_agent:
                    hier_result = await self._execute_hierarchical(
                        db, execution, tasks_with_agent, agents_map, task_map,
                        task_outputs, task_exec_map, completed_task_ids,
                        total_tokens, total_cost, completed_count, total_tasks,
                    )
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.total_tokens_used = hier_result["total_tokens"]
                    execution.total_cost_usd = hier_result["total_cost"]
                    execution.results = {"outputs": task_outputs, "mode": "hierarchical", **hier_result}
                    self._add_trace(execution, "crew.completed", data={"mode": "hierarchical"})
                    await db.commit()
                    return

                # Evaluator-Optimizer 模式：评估 + 优化循环
                if is_evaluator_optimizer and len(agents_map) >= 2 and tasks_with_agent:
                    eo_result = await self._execute_evaluator_optimizer(
                        db, execution, tasks_with_agent, agents_map, task_map,
                        task_outputs, task_exec_map, completed_task_ids,
                        total_tokens, total_cost, completed_count, total_tasks,
                    )
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.total_tokens_used = eo_result["total_tokens"]
                    execution.total_cost_usd = eo_result["total_cost"]
                    execution.results = {"outputs": task_outputs, "mode": "evaluator_optimizer", **eo_result}
                    self._add_trace(execution, "crew.completed", data={"mode": "evaluator_optimizer"})
                    await db.commit()
                    return

                # Prompt Chain 模式：链式顺序处理
                if is_prompt_chain:
                    pc_result = await self._execute_prompt_chain(
                        db, execution, tasks_with_agent, agents_map, task_map,
                        task_outputs, task_exec_map, completed_task_ids,
                        total_tokens, total_cost, completed_count, total_tasks,
                    )
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.total_tokens_used = pc_result["total_tokens"]
                    execution.total_cost_usd = pc_result["total_cost"]
                    execution.results = {"outputs": task_outputs, "mode": "prompt_chain", **pc_result}
                    self._add_trace(execution, "crew.completed", data={"mode": "prompt_chain"})
                    await db.commit()
                    return

                # Router 模式：分类 + 路由
                if is_router and len(agents_map) >= 2 and tasks_with_agent:
                    router_result = await self._execute_router(
                        db, execution, tasks_with_agent, agents_map, task_map,
                        task_outputs, task_exec_map, completed_task_ids,
                        total_tokens, total_cost, completed_count, total_tasks,
                    )
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.total_tokens_used = router_result["total_tokens"]
                    execution.total_cost_usd = router_result["total_cost"]
                    execution.results = {"outputs": task_outputs, "mode": "router", **router_result}
                    self._add_trace(execution, "crew.completed", data={"mode": "router"})
                    await db.commit()
                    return

                # Event Flow 模式：使用事件驱动编排器
                if is_event_flow:
                    event_flow_result = await self._execute_event_flow(
                        db, execution, crew, tasks_with_agent, agents_map, task_map,
                        task_outputs, task_exec_map, completed_task_ids,
                        total_tokens, total_cost, completed_count, total_tasks,
                    )
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.total_tokens_used = event_flow_result["total_tokens"]
                    execution.total_cost_usd = event_flow_result["total_cost"]
                    execution.results = {"outputs": task_outputs, "mode": "event_flow", **event_flow_result}
                    self._add_trace(execution, "crew.completed", data={"mode": "event_flow"})
                    await db.commit()
                    await event_publisher.publish_crew_completed(
                        execution_id=execution_id,
                        total_tokens=event_flow_result["total_tokens"],
                        total_cost=event_flow_result["total_cost"],
                        tasks_completed=event_flow_result["completed_count"],
                    )
                    return

                # Orchestrator-Workers 模式：Orchestrator 动态分配任务给 Worker
                if is_orchestrator and len(agents_map) >= 2 and tasks_with_agent:
                    orch_result = await self._execute_orchestrator_workers(
                        db, execution, tasks_with_agent, agents_map, task_map,
                        task_outputs, task_exec_map, completed_task_ids,
                        total_tokens, total_cost, completed_count, total_tasks,
                        checkpoint_manager, loop_outputs,
                    )
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.total_tokens_used = orch_result["total_tokens"]
                    execution.total_cost_usd = orch_result["total_cost"]
                    execution.results = {"outputs": task_outputs, "mode": "orchestrator", **orch_result}
                    self._add_trace(execution, "crew.completed", data={"mode": "orchestrator"})
                    await db.commit()
                    await event_publisher.publish_crew_completed(
                        execution_id=execution_id,
                        total_tokens=orch_result["total_tokens"],
                        total_cost=orch_result["total_cost"],
                        tasks_completed=orch_result["completed_count"],
                    )
                    return

                # 拓扑排序并分层
                levels = self._topological_levels(tasks_with_agent, task_map)

                if not levels and not loop_task_ids:
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.results = {"message": "没有可执行的任务", "total_tasks": len(all_tasks), "loop_results": loop_outputs}
                    self._add_trace(execution, "crew.completed", data={"reason": "无可用任务"})
                    await db.commit()
                    return
                elif not levels:
                    # 只有循环任务，全部已完成
                    execution.status = ExecutionStatus.COMPLETED
                    execution.completed_at = datetime.utcnow()
                    execution.total_tokens_used = total_tokens
                    execution.total_cost_usd = total_cost
                    execution.results = {
                        "outputs": task_outputs,
                        "total_tasks": total_tasks,
                        "completed_tasks": completed_count,
                        "loop_results": loop_outputs,
                    }
                    self._add_trace(execution, "crew.completed", data={
                        "total_tokens": total_tokens,
                        "total_cost": f"${total_cost:.4f}",
                        "tasks_completed": completed_count,
                    })
                    await db.commit()
                    return

                # 按层级执行
                for level_idx, level_tasks in enumerate(levels):
                    # 检查是否取消
                    if self._cancelled:
                        await self._cancel_remaining(db, execution, task_exec_map)
                        return

                    # 检查暂停信号（通过数据库暂停请求）
                    if await checkpoint_manager.check_pause_requested():
                        self._paused = True

                    # 如果已暂停：创建手动断点并退出
                    if self._paused:
                        await checkpoint_manager.create_manual_checkpoint(
                            completed_task_ids=list(completed_task_ids),
                            task_outputs=dict(task_outputs),
                            total_tokens=total_tokens,
                            total_cost_usd=total_cost,
                        )
                        await checkpoint_manager.complete_pause_request()
                        execution.status = ExecutionStatus.PAUSED
                        self._add_trace(execution, "crew.paused", data={
                            "completed_tasks": completed_count,
                            "total_tasks": total_tasks,
                        })
                        await db.commit()

                        await event_publisher.publish_progress(
                            execution_id=execution_id,
                            completed=completed_count,
                            total=total_tasks,
                        )
                        return

                    # 过滤已完成的任务（断点续传时跳过）
                    pending_tasks = [
                        t for t in level_tasks
                        if t.id not in completed_task_ids
                    ]

                    if not pending_tasks:
                        # 当前层级所有任务已完成，继续下一层
                        continue

                    if is_parallel and len(pending_tasks) > 1:
                        # 并行执行同一层级的任务（每个任务独立session）
                        self._add_trace(execution, "crew.parallel_start", data={
                            "level": level_idx + 1, "tasks": [t.name for t in pending_tasks]
                        })
                        await db.commit()

                        # 发布进度更新
                        await event_publisher.publish_progress(
                            execution_id=execution_id,
                            completed=completed_count,
                            total=total_tasks,
                        )

                        async def run_task_in_session(task):
                            async with get_db_session() as task_db:
                                te = task_exec_map[task.id]
                                # 从主session刷新te引用
                                te = await task_db.merge(te)
                                # I5: 不merge主session的execution对象，改用get避免并行冲突
                                execution_ref = await task_db.get(Execution, execution_id)
                                return await self._execute_task(
                                    task_db, execution_ref, task, agents_map, task_outputs, task_exec_map, te
                                )

                        results = await asyncio.gather(
                            *[run_task_in_session(task) for task in pending_tasks],
                            return_exceptions=True
                        )
                        # 刷新主session以获取最新状态
                        await db.refresh(execution)
                        for te in task_exec_map.values():
                            await db.refresh(te)

                        for task, result in zip(pending_tasks, results):
                            if isinstance(result, Exception):
                                await self._handle_task_failure(db, execution, task, task_exec_map, str(result), total_tokens, total_cost, total_tasks)
                                return
                            success, task_tokens, task_cost = result
                            if not success:
                                return
                            total_tokens += task_tokens
                            total_cost += task_cost
                            completed_count += 1
                            completed_task_ids.add(task.id)
                            await self._check_cost_alert(execution_id, total_cost)

                            # 实时广播费用/Token更新
                            await event_publisher.publish_cost_update(
                                execution_id=execution_id,
                                total_tokens=total_tokens,
                                total_cost=total_cost,
                            )

                            # 每个并行任务完成后自动创建断点
                            await checkpoint_manager.create_auto_checkpoint(
                                task_id=task.id,
                                task_name=task.name,
                                completed_task_ids=list(completed_task_ids),
                                task_outputs=dict(task_outputs),
                                total_tokens=total_tokens,
                                total_cost_usd=total_cost,
                            )
                    else:
                        # 顺序执行
                        for task in pending_tasks:
                            if self._cancelled:
                                await self._cancel_remaining(db, execution, task_exec_map)
                                return
                            success, task_tokens, task_cost = await self._execute_task(db, execution, task, agents_map, task_outputs, task_exec_map)
                            if not success:
                                return  # task failed
                            total_tokens += task_tokens
                            total_cost += task_cost
                            completed_count += 1
                            completed_task_ids.add(task.id)
                            await self._check_cost_alert(execution_id, total_cost)

                            # 实时广播费用/Token更新
                            await event_publisher.publish_cost_update(
                                execution_id=execution_id,
                                total_tokens=total_tokens,
                                total_cost=total_cost,
                            )

                            # 每个任务完成后自动创建断点
                            await checkpoint_manager.create_auto_checkpoint(
                                task_id=task.id,
                                task_name=task.name,
                                completed_task_ids=list(completed_task_ids),
                                task_outputs=dict(task_outputs),
                                total_tokens=total_tokens,
                                total_cost_usd=total_cost,
                            )

                # 所有任务完成
                execution.status = ExecutionStatus.COMPLETED
                execution.completed_at = datetime.utcnow()
                execution.total_tokens_used = total_tokens
                execution.total_cost_usd = total_cost
                execution.results = {
                    "outputs": task_outputs,
                    "total_tasks": total_tasks,
                    "completed_tasks": completed_count,
                    "loop_results": loop_outputs,
                }
                self._add_trace(execution, "crew.completed", data={
                    "total_tokens": total_tokens,
                    "total_cost": f"${total_cost:.4f}",
                    "tasks_completed": completed_count,
                })
                await db.commit()

                # 发布工作流完成事件
                await event_publisher.publish_crew_completed(
                    execution_id=execution_id,
                    total_tokens=total_tokens,
                    total_cost=total_cost,
                    tasks_completed=completed_count,
                )

                logger.info(f"Execution {execution_id} completed successfully")

                # 分层记忆：将执行结果写入工作空间的 .fugue/ 目录
                memory_result = await self._persist_memory_to_workspace(
                    db, execution, crew, agents_map, task_outputs, total_tokens, total_cost
                )
                if memory_result:
                    execution.results["memory_persisted"] = memory_result

                # 将 worktree 中的输出文件递归复制回用户工作空间（防止丢失）
                if getattr(self, '_worktree_path', None) and getattr(crew, 'workspace_dir', None):
                    try:
                        import shutil, os as _os
                        wt_path = self._worktree_path
                        user_ws = _os.path.expanduser(crew.workspace_dir)
                        if _os.path.isdir(wt_path) and wt_path != user_ws:
                            copied = 0
                            for root, dirs, files in _os.walk(wt_path):
                                for fname in files:
                                    src = _os.path.join(root, fname)
                                    rel = _os.path.relpath(src, wt_path)
                                    dst = _os.path.join(user_ws, rel)
                                    if not _os.path.exists(dst):
                                        _os.makedirs(_os.path.dirname(dst), exist_ok=True)
                                        shutil.copy2(src, dst)
                                        copied += 1
                            if copied > 0:
                                logger.info(f"[EXECUTOR] Copied {copied} output file(s) from worktree to {user_ws}")
                    except Exception as copy_err:
                        logger.warning(f"[EXECUTOR] Output copy failed: {copy_err}")

                # 清理 Git Worktree
                if getattr(self, '_worktree_path', None):
                    try:
                        from app.services.worktree_manager import get_worktree_manager
                        wt_mgr = get_worktree_manager()
                        await wt_mgr.remove_worktree(self._worktree_path)
                        logger.info(f"[EXECUTOR] Worktree removed: {self._worktree_path}")
                    except Exception as wt_err:
                        logger.warning(f"[EXECUTOR] Worktree cleanup failed: {wt_err}")
                    self._worktree_path = None

                # 清除工作空间设置
                from app.engine.tools import clear_workspace_dir
                clear_workspace_dir()

        except Exception as e:
            logger.error(f"Execution {execution_id} failed with exception: {e}\n{traceback.format_exc()}")
            # 异常路径也复制输出文件
            if getattr(self, '_worktree_path', None) and getattr(crew, 'workspace_dir', None):
                try:
                    import shutil, os as _os
                    copied = 0
                    for root, dirs, files in _os.walk(self._worktree_path):
                        for fname in files:
                            rel = _os.path.relpath(_os.path.join(root, fname), self._worktree_path)
                            dst = _os.path.join(_os.path.expanduser(crew.workspace_dir), rel)
                            if not _os.path.exists(dst):
                                _os.makedirs(_os.path.dirname(dst), exist_ok=True)
                                shutil.copy2(_os.path.join(root, fname), dst)
                                copied += 1
                    if copied > 0:
                        logger.info(f"[EXECUTOR] Error path: copied {copied} files to workspace")
                except Exception as e:
                    logger.warning(f"Error path file copy failed: {e}")
            # 清理 Git Worktree（异常路径）
            if getattr(self, '_worktree_path', None):
                try:
                    from app.services.worktree_manager import get_worktree_manager
                    wt_mgr = get_worktree_manager()
                    await wt_mgr.remove_worktree(self._worktree_path)
                    logger.info(f"[EXECUTOR] Worktree removed after error: {self._worktree_path}")
                except Exception as wt_err:
                    logger.warning(f"[EXECUTOR] Worktree cleanup failed: {wt_err}")
                self._worktree_path = None
            # 清除工作空间设置（异常路径也要清理，防止污染后续执行）
            try:
                from app.engine.tools import clear_workspace_dir
                clear_workspace_dir()
            except Exception as e:
                logger.warning(f"Error path workspace cleanup failed: {e}")
            # 最后的防线：确保异常情况下也能更新数据库状态
            try:
                async with get_db_session() as db:
                    execution = await db.get(Execution, execution_id)
                    if execution and execution.status in [ExecutionStatus.PENDING, ExecutionStatus.RUNNING]:
                        execution.status = ExecutionStatus.FAILED
                        execution.completed_at = datetime.utcnow()
                        execution.error_log = f"执行异常: {str(e)[:500]}"
                        self._add_trace(execution, "crew.failed", data={"error": str(e)[:200]})
                        await db.commit()
                        logger.info(f"Marked execution {execution_id} as FAILED after exception")

                        # 发布工作流失败事件
                        await event_publisher.publish_crew_failed(
                            execution_id=execution_id,
                            error=str(e),
                        )
            except Exception as inner_e:
                logger.critical(f"Failed to mark execution {execution_id} as FAILED: {inner_e}")
        finally:
            # 确保所有退出路径（包括 return）都清理 workspace
            try:
                from app.engine.tools import clear_workspace_dir
                clear_workspace_dir()
            except Exception as e:
                logger.debug(f"Workspace cleanup in finally failed (non-critical): {e}")
            if getattr(self, '_worktree_path', None):
                try:
                    from app.services.worktree_manager import get_worktree_manager
                    import asyncio as _aio
                    # 使用 create_task 并等待最多3秒完成清理
                    cleanup_task = _aio.create_task(
                        get_worktree_manager().remove_worktree(self._worktree_path)
                    )
                    try:
                        await _aio.wait_for(cleanup_task, timeout=3.0)
                    except _aio.TimeoutError:
                        cleanup_task.cancel()
                except Exception as e:
                    logger.debug(f"Worktree cleanup in finally failed (non-critical): {e}")
                self._worktree_path = None

    # --- 附件注入 ---
    MAX_FULL_INJECT = 10 * 1024 * 1024  # 10MB
    TRUNCATE_PREVIEW = 100 * 1024        # 100KB

    def _extract_pdf_text(self, file_path: str) -> str:
        """从 PDF 文件提取文本内容"""
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages[:100]):  # 最多处理 100 页
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- 第 {i+1} 页 ---\n{page_text}")
        if not text_parts:
            raise ValueError("PDF 中未提取到文本内容（可能是扫描型 PDF，暂不支持 OCR）")
        return "\n\n".join(text_parts)

    async def _inject_attachments(self, task, context_parts: list) -> list:
        """将预设文件内容注入到任务上下文"""
        attachments = []
        if hasattr(task, 'config') and task.config:
            attachments = task.config.get("attachments", [])
        if not attachments:
            return context_parts

        from app.engine.local_fs_client import read_file

        file_contents = []
        for att in attachments:
            try:
                file_path = att["path"]
                mime = att.get("mime_type", "")
                name = att.get("name", "unknown")
                ext = os.path.splitext(file_path)[1].lower()

                # PDF 文件：提取文本
                if mime == "application/pdf" or ext == ".pdf":
                    content = await asyncio.to_thread(self._extract_pdf_text, file_path)
                # DOCX 文件：用 python-docx 提取
                elif ext == ".docx" or "wordprocessingml" in mime:
                    content = await asyncio.to_thread(self._extract_docx_text, file_path)
                # XLSX 文件：用 openpyxl 提取
                elif ext in (".xlsx", ".xls") or "spreadsheet" in mime:
                    content = await asyncio.to_thread(self._extract_xlsx_text, file_path)
                # 图片文件：记录信息但不注入内容
                elif mime.startswith("image/"):
                    file_contents.append(
                        f"[附件: {name}] 图片文件（{att.get('size', 0) // 1024}KB）。"
                        f"路径: {file_path}。如需分析请使用 image_info 工具。"
                    )
                    continue
                # 文本文件：通过 Tauri 代理读取
                else:
                    content = await read_file(file_path)

                if len(content) > self.MAX_FULL_INJECT:
                    content = content[:self.TRUNCATE_PREVIEW] + \
                        f"\n... [文件过大已截断，仅显示前 100KB，共 {att.get('size', '?')} 字节。"
                file_contents.append(f"[附件: {name}]\n```\n{content}\n```")
            except FileNotFoundError:
                file_contents.append(
                    f"[附件: {att.get('name', 'unknown')}] 读取失败: 文件不存在。"
                    f"原路径: {att['path']}"
                )
            except Exception as e:
                file_contents.append(f"[附件: {att.get('name', 'unknown')}] 读取失败: {e}")

        if file_contents:
            context_parts.insert(0, "以下是用户上传的参考文件（已自动读取，无需再用工具读取）：\n\n" + "\n\n".join(file_contents))

        return context_parts

    def _extract_docx_text(self, file_path: str) -> str:
        """从 DOCX 文件提取文本内容"""
        from docx import Document
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            parts.append("[表格]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts) if parts else "(空文档)"

    def _extract_xlsx_text(self, file_path: str) -> str:
        """从 XLSX 文件提取文本内容"""
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"[工作表: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts) if parts else "(空工作簿)"

    async def _execute_task(
        self, db: AsyncSession, execution: Execution,
        task: Task, agents_map: Dict[str, Agent],
        task_outputs: Dict[str, str],
        task_exec_map: Dict[str, TaskExecution],
        te_override: TaskExecution = None,
    ):
        """执行单个任务，返回 (success: bool, tokens_used: int, cost_usd: float)"""
        # 工具调用循环本地累计 token 和费用
        local_total_tokens = 0
        local_total_cost = 0.0

        te = te_override or task_exec_map[task.id]
        agent = agents_map.get(task.agent_id)

        if not agent:
            te.status = TaskExecutionStatus.SKIPPED
            te.error_message = "未分配Agent"
            self._add_trace(execution, "task.skipped", task_name=task.name, data={"reason": "未分配Agent"})
            await db.commit()
            return True, 0, 0.0

        te.status = TaskExecutionStatus.RUNNING
        te.started_at = datetime.utcnow()
        self._add_trace(execution, "task.started", agent_name=agent.name, task_name=task.name)
        await db.commit()

        # 子工作流执行（嵌套 Crew）
        if task.sub_crew_id:
            return await self._execute_sub_crew(db, execution, task, te)

        # 发布任务开始事件
        await event_publisher.publish_task_started(
            execution_id=self.execution_id,
            task_name=task.name,
            agent_name=agent.name,
        )

        # 成本预算检查
        cost_budget = getattr(self, '_cost_budget', None)
        if cost_budget and execution.total_cost_usd and execution.total_cost_usd >= cost_budget:
            te.status = TaskExecutionStatus.FAILED
            te.error_message = f"已达到成本预算上限 ${cost_budget:.2f}，任务中止"
            te.completed_at = datetime.utcnow()
            self._add_trace(execution, "task.failed", agent_name=agent.name, task_name=task.name,
                           data={"error": te.error_message})
            await db.commit()
            return False, 0, 0.0

        # 重试逻辑
        max_retries = task.max_retries or 1
        _task_start = time.monotonic()
        for attempt in range(max_retries):
            try:
                # 构建上下文
                context_parts = []
                for dep_id in (task.context_task_ids or []):
                    if dep_id in task_outputs:
                        context_parts.append(f"[依赖任务输出]:\n{task_outputs[dep_id]}")

                # 注入记忆上下文（短期记忆 + RAG 知识库检索）— 超时保护
                memory_context = ""
                try:
                    memory_context = await asyncio.wait_for(
                        self._build_memory_context(db, agent, task, execution),
                        timeout=15,  # 最多等15秒，超时跳过
                    )
                except asyncio.TimeoutError:
                    logger.warning(f"[EXECUTOR] Memory context build timed out (15s), skipping for task {task.name}")
                except Exception as mem_err:
                    logger.warning(f"[EXECUTOR] Memory context build failed: {mem_err}")
                if memory_context:
                    context_parts.append(memory_context)
                logger.info(f"[PERF] Memory context: {time.monotonic() - _task_start:.1f}s")

                # 注入附件内容
                context_parts = await self._inject_attachments(task, context_parts)

                messages = self._build_messages(agent, task, context_parts)
                logger.info(f"[PERF] Messages built: {time.monotonic() - _task_start:.1f}s")
                self._add_trace(execution, "agent.thinking", agent_name=agent.name, data={"step": f"调用LLM (尝试 {attempt+1}/{max_retries})"})

                # 发布Agent思考事件
                await event_publisher.publish_agent_thinking(
                    execution_id=self.execution_id,
                    agent_name=agent.name,
                    thought=f"准备调用LLM (尝试 {attempt+1}/{max_retries})",
                    step="llm_call",
                )

                provider_key = self.llm_api_keys.get(agent.llm_provider)
                provider_url = self.llm_base_urls.get(agent.llm_provider)
                llm = get_llm_provider(
                    agent.llm_provider,
                    api_key=provider_key,
                    base_url=provider_url,
                    all_keys=self.llm_api_keys,
                    all_base_urls=self.llm_base_urls,
                )

                # 检测是否回退到了 MockProvider
                from app.engine.llm_provider import MockProvider
                if isinstance(llm, MockProvider):
                    warning_msg = (
                        f"Agent [{agent.name}] 的 LLM Provider '{agent.llm_provider}' "
                        f"未找到有效 API Key，已回退到演示模式（Mock）。"
                        f"输出为模拟数据，请在「设置」中配置 API Key。"
                    )
                    logger.warning(warning_msg)
                    self._add_trace(execution, "agent.warning", agent_name=agent.name, data={
                        "step": warning_msg,
                    })

                timeout = task.timeout_seconds or 300

                # 构建工具 schema — 始终注入所有可用工具
                agent_tool_names = agent.tools_config or []
                tool_schemas = []
                is_anthropic = agent.llm_provider == "anthropic"

                # 1) 内置工具（按 agent 配置筛选）
                if agent_tool_names:
                    tool_schemas = get_anthropic_tools(agent_tool_names) if is_anthropic else get_openai_tools(agent_tool_names)

                # 2) MCP 工具
                from app.engine.mcp_adapter import get_mcp_adapter
                mcp_schemas = get_mcp_adapter().get_tool_schemas()
                if mcp_schemas:
                    tool_schemas = (tool_schemas or []) + mcp_schemas

                # 3) 所有插件工具（LocalFS + EnhancedTools 等）— 始终注入
                provider_fmt = "anthropic" if is_anthropic else "openai"
                plugin_schemas = get_plugin_tool_schemas(provider_fmt)
                if plugin_schemas:
                    tool_schemas = (tool_schemas or []) + plugin_schemas

                # 4) 如果 agent 无 tools_config 且插件也无 schema，兜底获取内置工具
                if not agent_tool_names and not tool_schemas:
                    _default_tools = ["web_search", "file_read", "file_write", "code_execute", "api_call", "text_analysis"]
                    builtin_schemas = get_anthropic_tools(_default_tools) if is_anthropic else get_openai_tools(_default_tools)
                    if builtin_schemas:
                        tool_schemas = builtin_schemas

                logger.info(
                    f"[EXECUTOR] Agent '{agent.name}' tools: "
                    f"agent_config={len(agent_tool_names)}, "
                    f"plugin={len(plugin_schemas) if plugin_schemas else 0}, "
                    f"total_schemas={len(tool_schemas or [])}"
                )
                logger.info(f"[PERF] Tool schemas ready: {time.monotonic() - _task_start:.1f}s")

                # 多轮工具调用循环
                all_tool_calls = []
                max_tool_rounds = self.max_turns

                # 自动降级：预算紧张时切换到更便宜的模型
                from app.engine.llm_provider import select_degraded_model
                effective_model = select_degraded_model(
                    agent.llm_model,
                    budget_remaining=(getattr(self, '_cost_budget', None) or 0) - (execution.total_cost_usd or 0),
                    budget_total=getattr(self, '_cost_budget', None),
                )

                for tool_round in range(max_tool_rounds):
                    # 流式 LLM 调用 — 实时推送思考过程
                    llm_response: LLMResponse = None
                    stream_buffer = ""

                    try:
                        stream_gen = llm.chat_stream(
                            messages=messages,
                            model=effective_model,
                            temperature=agent.temperature,
                            max_tokens=agent.max_tokens,
                            tools=tool_schemas if tool_schemas else None,
                        )
                        async for event in stream_gen:
                            if event is None:
                                continue
                            etype = getattr(event, 'event_type', None)
                            if etype == "text_delta":
                                text_chunk = event.data or ""
                                if text_chunk:
                                    stream_buffer += text_chunk
                                    if len(stream_buffer) % 200 < len(text_chunk):
                                        await event_publisher.publish_agent_thinking(
                                            execution_id=self.execution_id,
                                            agent_name=agent.name,
                                            thought=stream_buffer[-500:],
                                            step=f"生成中 (轮次 {tool_round+1})",
                                        )
                            elif etype == "tool_call_start":
                                tc_info = event.data or {}
                                await event_publisher.publish_agent_thinking(
                                    execution_id=self.execution_id,
                                    agent_name=agent.name,
                                    thought=f"正在调用工具: {tc_info.get('name', '?')}",
                                    step="tool_call",
                                )
                            elif etype == "done":
                                llm_response = event.data

                    except asyncio.TimeoutError:
                        logger.warning(f"[EXECUTOR] LLM streaming timeout ({timeout}s)")
                        raise
                    except Exception as stream_err:
                        # 流式失败 → 回退到非流式
                        logger.warning(f"[EXECUTOR] Streaming error: {stream_err}, falling back to non-streaming")
                        llm_response = await asyncio.wait_for(
                            llm.chat(
                                messages=messages, model=effective_model,
                                temperature=agent.temperature, max_tokens=agent.max_tokens,
                                tools=tool_schemas if tool_schemas else None,
                            ),
                            timeout=timeout,
                        )

                    if llm_response is None:
                        raise Exception("LLM 调用未返回响应")

                    # LLM 调用成功 → 更新健康状态
                    from app.engine.llm_provider import update_provider_health
                    await update_provider_health(agent.llm_provider, success=True)

                    # 累计 token 和费用
                    local_total_tokens += llm_response.tokens_used
                    local_total_cost += llm_response.cost_usd

                    # 推送模型的推理文本（如果有的话）
                    model_text = stream_buffer or (llm_response.content or "")
                    model_text = re.sub(r'```tool_call\s*\n.*?\n```', '', model_text, flags=re.DOTALL).strip()
                    if model_text:
                        await event_publisher.publish_agent_thinking(
                            execution_id=self.execution_id,
                            agent_name=agent.name,
                            thought=model_text[:2000],
                            step="reasoning",
                        )

                    # 为即将执行的工具生成推理描述（即使模型没有输出文本）
                    if llm_response.tool_calls:
                        for tc in llm_response.tool_calls:
                            tool_desc = self._describe_tool_call(tc.name, tc.arguments)
                            await event_publisher.publish_agent_thinking(
                                execution_id=self.execution_id,
                                agent_name=agent.name,
                                thought=tool_desc,
                                step="reasoning",
                            )

                    # 无工具调用 → 检查 Agent Command API（goto/update/resume）
                    if not llm_response.tool_calls:
                        cmd_name, cmd_args = self._parse_agent_command(llm_response.content or "")
                        if cmd_name:
                            logger.info(f"[EXECUTOR] Agent command detected: {cmd_name}")
                            cmd_result = await self._handle_agent_command(
                                db, execution, cmd_name, cmd_args or {},
                                agent, task, agents_map, task_outputs, task_exec_map,
                            )
                            if cmd_result["handled"] and cmd_name == "goto":
                                # goto 成功后，重新执行当前任务（由新 Agent 接手）
                                self._add_trace(execution, "agent.command", agent_name=agent.name, data=cmd_result)
                                te.status = TaskExecutionStatus.RETRYING
                                await db.commit()
                                return False, 0, 0.0  # 标记为需要重试

                    # 无工具调用 → 检查文本工具调用协议（备用）
                    if not llm_response.tool_calls:
                        text_tool_calls = self._parse_text_tool_calls(llm_response.content or "")
                        if text_tool_calls:
                            logger.info(f"[EXECUTOR] Parsed {len(text_tool_calls)} text-based tool calls from response")
                            llm_response.tool_calls = text_tool_calls
                            clean_content = re.sub(r'```tool_call\s*\n.*?\n```', '', llm_response.content, flags=re.DOTALL).strip()
                            if clean_content:
                                llm_response.content = clean_content
                        else:
                            final_content = llm_response.content
                            break

                    # LLM 请求了工具调用
                    self._add_trace(execution, "agent.thinking", agent_name=agent.name, data={
                        "step": f"工具调用轮次 {tool_round + 1}: {', '.join(tc.name for tc in llm_response.tool_calls)}",
                    })

                    # 执行每个工具
                    is_anthropic = agent.llm_provider == "anthropic"

                    for tc in llm_response.tool_calls:
                        await event_publisher.publish_agent_tool_call(
                            execution_id=self.execution_id,
                            agent_name=agent.name,
                            tool_name=tc.name,
                            tool_input=tc.arguments,
                        )

                        # 审批检查：根据 crew 的 approval_mode 决定是否需要人工确认
                        from app.engine.tools import ToolResult
                        from app.models.crew import Crew
                        _crew = await db.get(Crew, execution.crew_id)
                        crew_approval_mode = getattr(_crew, 'approval_mode', 'semi_auto') or 'semi_auto'
                        if crew_approval_mode != 'full_auto':
                            from app.services.approval_manager import get_approval_manager, ApprovalMode
                            approval_mgr = get_approval_manager()
                            mode = ApprovalMode(crew_approval_mode)
                            if approval_mgr.requires_approval(mode=mode, tool_name=tc.name):
                                # 创建审批请求并等待
                                approval_req = await approval_mgr.create_approval_request(
                                    execution_id=self.execution_id,
                                    tool_name=tc.name,
                                    tool_args=tc.arguments,
                                )
                                # 通知前端需要审批
                                _tool_cn = {
                                    "file_write": "写入文件", "file_read": "读取文件",
                                    "code_execute": "执行代码", "shell_execute": "执行命令",
                                    "database_query": "数据库查询", "api_call": "调用接口",
                                    "web_search": "网络搜索", "image_generation": "生成图片",
                                    "text_analysis": "文本分析", "fs_write": "写入文件",
                                    "fs_read": "读取文件", "fs_list": "浏览目录",
                                    "fs_info": "查看文件信息", "fs_mkdir": "创建文件夹",
                                    "fs_delete": "删除文件", "fs_move": "移动文件",
                                    "fs_copy": "复制文件", "doc_read": "读取文档",
                                    "doc_write": "创建文档", "doc_edit": "编辑文档",
                                    "doc_list": "浏览文档列表",
                                    "docx_read": "读取文档", "docx_create": "创建文档",
                                    "docx_edit": "编辑文档", "docx_list": "浏览文档列表",
                                    "docx_delete": "删除文档",
                                    "calculator": "计算",
                                    "execute_workflow": "运行工作流", "send_message": "发送消息",
                                    "read_page": "读取网页",
                                }
                                _cn_name = _tool_cn.get(tc.name, tc.name)
                                _mode_cn = {"safe": "限制权限", "semi_auto": "默认权限", "full_auto": "完全权限"}
                                await event_publisher.publish_warning(
                                    execution_id=self.execution_id,
                                    message=f"{_cn_name}操作需要您审批（{_mode_cn.get(crew_approval_mode, crew_approval_mode)}模式）",
                                    extra={
                                        "approval_request_id": approval_req["request_id"],
                                        "tool_name": tc.name,
                                        "risk_level": approval_req["risk_level"],
                                        "tool_args": tc.arguments,
                                    },
                                )
                                # 更新执行状态
                                execution.status = ExecutionStatus.WAITING_REVIEW
                                await db.commit()
                                # 等待审批结果
                                approval_result = await approval_mgr.wait_for_approval(
                                    approval_req["request_id"], timeout=600
                                )
                                if approval_result["status"] != "approved":
                                    tool_result = ToolResult(success=False, output="", error=f"工具调用被拒绝: {approval_result.get('reject_reason', '用户拒绝')}")
                                    tc.result = tool_result.output
                                    all_tool_calls.append({
                                        "timestamp": datetime.utcnow().isoformat(),
                                        "tool_name": tc.name,
                                        "input": tc.arguments,
                                        "output": "",
                                        "duration_ms": 0,
                                        "error": tool_result.error,
                                    })
                                    continue
                                # 审批通过，恢复执行状态
                                execution.status = ExecutionStatus.RUNNING
                                await db.commit()

                        # MCP 适配器（沙箱和 MCP 路由共用）
                        from app.engine.mcp_adapter import get_mcp_adapter
                        from app.engine.tools import ToolResult as TR
                        mcp_adapter = get_mcp_adapter()

                        # 沙箱包装：对高危工具（code_execute, shell_execute）使用沙箱执行
                        SANDBOX_WRAPPED_TOOLS = {"code_execute", "shell_execute"}
                        if tc.name in SANDBOX_WRAPPED_TOOLS:
                            from app.engine.sandbox import get_sandbox_manager, SandboxConfig
                            _sb_mgr = get_sandbox_manager()
                            _sb_config = SandboxConfig(
                                enable_filesystem_isolation=True,
                                enable_network_isolation=False,
                                max_execution_time=120,
                                allowed_paths=[
                                    str(Path.home() / ".fugue"),
                                    str(Path.home() / "Documents"),
                                ],
                            )
                            _command = tc.arguments.get("command") or tc.arguments.get("code") or ""
                            _workspace = getattr(self, '_worktree_path', None) or str(Path.home() / ".fugue" / "workspace")
                            _sb_result = await _sb_mgr.execute_in_sandbox(
                                command=_command,
                                workspace=_workspace,
                                config=_sb_config,
                            )
                            tool_result = TR(
                                tool_call_id=tc.id, tool_name=tc.name,
                                output=_sb_result.get("output", ""),
                                success=_sb_result.get("success", False),
                                error=_sb_result.get("error"),
                            )
                        # MCP 工具路由
                        elif mcp_adapter.is_mcp_tool(tc.name):
                            server_id, original_name = mcp_adapter.parse_mcp_tool_name(tc.name)
                            mcp_result = await mcp_adapter.call_tool(server_id, original_name, tc.arguments)
                            tool_result = TR(
                                success=mcp_result.get("success", False),
                                output=mcp_result.get("output", ""),
                                error=mcp_result.get("error"),
                            )
                        else:
                            tool_result = await execute_tool(tc.name, tc.arguments, tc.id)
                        tc.result = tool_result.output

                        all_tool_calls.append({
                            "timestamp": datetime.utcnow().isoformat(),
                            "tool_name": tc.name,
                            "input": tc.arguments,
                            "output": tool_result.output,
                            "duration_ms": tool_result.duration_ms,
                            "error": tool_result.error,
                        })

                        # 推送工具结果作为思考事件（让前端实时显示）
                        tool_summary = (tool_result.output or "")[:300]
                        if tool_result.error:
                            tool_summary = f"错误: {tool_result.error[:200]}"
                        await event_publisher.publish_agent_thinking(
                            execution_id=self.execution_id,
                            agent_name=agent.name,
                            thought=f"[{tc.name}] {tool_summary}",
                            step="tool_result",
                        )

                        self._add_trace(execution, "agent.tool_call", agent_name=agent.name, task_name=task.name, data={
                            "tool": tc.name, "success": tool_result.success, "duration_ms": tool_result.duration_ms,
                        })

                    # 将工具调用和结果追加到消息历史
                    if is_anthropic:
                        # Anthropic: assistant 消息包含 tool_use blocks
                        assistant_blocks = []
                        if llm_response.content:
                            assistant_blocks.append({"type": "text", "text": llm_response.content})
                        for tc in llm_response.tool_calls:
                            assistant_blocks.append({"type": "tool_use", "id": tc.id, "name": tc.name, "input": tc.arguments})
                        messages.append({"role": "assistant", "content": assistant_blocks})
                        # 工具结果作为 user 消息
                        tool_result_blocks = []
                        for tc in llm_response.tool_calls:
                            tool_result_blocks.append({
                                "type": "tool_result",
                                "tool_use_id": tc.id,
                                "content": (tc.result or "")[:8000],
                            })
                        messages.append({"role": "user", "content": tool_result_blocks})
                    else:
                        # OpenAI 格式
                        assistant_msg: Dict[str, Any] = {"role": "assistant", "content": llm_response.content or ""}
                        assistant_msg["tool_calls"] = [
                            {"id": tc.id, "type": "function", "function": {"name": tc.name, "arguments": json.dumps(tc.arguments, ensure_ascii=False)}}
                            for tc in llm_response.tool_calls
                        ]
                        messages.append(assistant_msg)
                        for tc in llm_response.tool_calls:
                            # 截断过长的工具结果，避免超出 API 上下文限制
                            result_content = (tc.result or "")[:8000]
                            messages.append({"role": "tool", "tool_call_id": tc.id, "content": result_content})

                    # 强制推理：单独发起一次无工具的 LLM 调用，让模型用自然语言解释
                    if tool_round < max_tool_rounds - 1 and llm_response.tool_calls:
                        try:
                            explain_messages = messages + [{
                                "role": "user",
                                "content": "请用2-3句中文自然语言，简要说明你刚才做了什么、结果如何、接下来打算做什么。不要调用任何工具，只输出文字。"
                            }]
                            explain_resp = await asyncio.wait_for(
                                llm.chat(messages=explain_messages, model=effective_model,
                                         temperature=0.7, max_tokens=200, tools=None),
                                timeout=30,
                            )
                            explain_text = (explain_resp.content or "").strip()
                            if explain_text:
                                local_total_tokens += explain_resp.tokens_used
                                local_total_cost += explain_resp.cost_usd
                                await event_publisher.publish_agent_thinking(
                                    execution_id=self.execution_id,
                                    agent_name=agent.name,
                                    thought=explain_text,
                                    step="reasoning",
                                )
                                # 将推理文本追加到消息历史，保持上下文连贯
                                messages.append({"role": "assistant", "content": explain_text})
                        except Exception as explain_err:
                            logger.debug(f"[EXECUTOR] Reasoning call failed (non-critical): {explain_err}")
                else:
                    # 工具循环达到上限 — 强制请求最终回复
                    logger.warning(f"[EXECUTOR] Agent '{agent.name}' hit max tool rounds ({max_tool_rounds}), requesting final answer")
                    try:
                        final_prompt_msg = {"role": "user", "content": "请立即基于以上所有工具调用的结果，给出完整的最终答案。不要再调用任何工具。"}
                        messages.append(final_prompt_msg)
                        final_resp = await asyncio.wait_for(
                            llm.chat(messages=messages, model=effective_model, temperature=agent.temperature, max_tokens=agent.max_tokens, tools=None),
                            timeout=timeout,
                        )
                        final_content = final_resp.content or ""
                        local_total_tokens += final_resp.tokens_used
                        local_total_cost += final_resp.cost_usd
                    except Exception as final_err:
                        logger.error(f"[EXECUTOR] Final answer request failed: {final_err}")
                        final_content = ""

                    if not final_content:
                        # 汇总工具调用结果作为兜底输出
                        tool_summary = []
                        for tc_entry in all_tool_calls[-5:]:
                            tn = tc_entry.get("tool_name", "?")
                            to = str(tc_entry.get("output", ""))[:300]
                            tool_summary.append(f"[{tn}] {to}")
                        final_content = f"[工具调用已达上限，以下为最近的工具结果]\n\n" + "\n\n".join(tool_summary) if tool_summary else "(工具调用轮次超限，未获得有效结果)"

                # 成功
                te.status = TaskExecutionStatus.COMPLETED
                te.completed_at = datetime.utcnow()
                te.output = final_content
                te.tokens_used = local_total_tokens
                te.cost_usd = local_total_cost
                # thoughts 包含所有工具调用记录 + 最终内容摘要
                thought_entries = []
                for tc_entry in all_tool_calls:
                    thought_entries.append({
                        "timestamp": tc_entry.get("timestamp", ""),
                        "type": "tool_use",
                        "content": f"[{tc_entry.get('tool_name','')}] {str(tc_entry.get('output',''))[:200]}",
                    })
                thought_entries.append({
                    "timestamp": datetime.utcnow().isoformat(),
                    "type": "reasoning",
                    "content": (final_content or "")[:500],
                })
                te.thoughts = thought_entries
                te.tool_calls = all_tool_calls
                flag_modified(te, 'thoughts')
                flag_modified(te, 'tool_calls')
                task_outputs[task.id] = final_content or ""
                self._add_trace(execution, "task.completed", agent_name=agent.name, task_name=task.name, data={
                    "tokens": local_total_tokens,
                    "cost": f"${local_total_cost:.4f}",
                    "tool_calls": len(all_tool_calls),
                })
                await db.commit()

                # P1-2: 任务完成后保存结论到 Agent 长期记忆
                if final_content and agent:
                    try:
                        from app.services.memory_service import MemoryService
                        mem_svc = MemoryService(db)
                        await mem_svc.save_memory(
                            agent_id=str(agent.id),
                            content=final_content[:2000],
                            memory_type="conclusion",
                            execution_id=str(execution.id),
                        )
                    except Exception as mem_err:
                        logger.warning(f"[EXECUTOR] Memory save failed: {mem_err}")

                # P1-1: 任务完成后自动索引输出到知识库
                if getattr(self, '_memory_config', None) and getattr(self._memory_config, 'auto_index_on_complete', True):
                    try:
                        from app.services.vector_store import get_vector_store
                        import uuid as _uuid
                        vs = get_vector_store()
                        if vs and final_content:
                            # 查找 agent 关联的知识库
                            from app.models.memory import AgentKnowledgeMapping
                            _mappings = await db.execute(
                                select(AgentKnowledgeMapping).where(AgentKnowledgeMapping.agent_id == agent.id)
                            )
                            _mapping = _mappings.scalars().first()
                            if _mapping:
                                await vs.add_documents(
                                    knowledge_base_id=str(_mapping.knowledge_base_id),
                                    documents=[{
                                        "id": str(_uuid.uuid4()),
                                        "content": final_content[:5000],
                                        "metadata": {
                                            "task_name": task.name,
                                            "agent_id": str(agent.id),
                                            "execution_id": str(execution.id),
                                            "type": "task_output",
                                        },
                                    }],
                                )
                                logger.info(f"[EXECUTOR] Auto-indexed task output to KB: {task.name}")
                    except Exception as idx_err:
                        logger.warning(f"[EXECUTOR] Auto-index failed: {idx_err}")

                # 发布任务完成事件
                await event_publisher.publish_task_completed(
                    execution_id=self.execution_id,
                    task_name=task.name,
                    agent_name=agent.name,
                    tokens_used=local_total_tokens,
                    cost_usd=local_total_cost,
                )

                return True, local_total_tokens, local_total_cost

            except asyncio.TimeoutError:
                te.retry_count = attempt + 1
                if attempt < max_retries - 1:
                    te.status = TaskExecutionStatus.RETRYING
                    self._add_trace(execution, "task.retrying", agent_name=agent.name, task_name=task.name, data={
                        "reason": f"超时({timeout}s)，重试 {attempt+2}/{max_retries}"
                    })
                    await db.commit()
                    await asyncio.sleep(2 ** attempt)
                else:
                    te.status = TaskExecutionStatus.FAILED
                    te.completed_at = datetime.utcnow()
                    te.error_message = f"任务超时({timeout}s)，已重试{max_retries}次"
                    self._add_trace(execution, "task.failed", agent_name=agent.name, task_name=task.name,
                                    data={"error": te.error_message})
                    await db.commit()

                    # 发布任务失败事件
                    await event_publisher.publish_task_failed(
                        execution_id=self.execution_id,
                        task_name=task.name,
                        agent_name=agent.name,
                        error=te.error_message,
                    )

                    return False, local_total_tokens, local_total_cost

            except Exception as e:
                # LLM 调用失败 → 更新健康状态
                from app.engine.llm_provider import update_provider_health
                await update_provider_health(agent.llm_provider, success=False)

                te.retry_count = attempt + 1
                if attempt < max_retries - 1:
                    te.status = TaskExecutionStatus.RETRYING
                    self._add_trace(execution, "task.retrying", agent_name=agent.name, task_name=task.name, data={
                        "reason": f"错误: {str(e)[:100]}，重试 {attempt+2}/{max_retries}"
                    })
                    await db.commit()
                    await asyncio.sleep(2 ** attempt)
                else:
                    te.status = TaskExecutionStatus.FAILED
                    te.completed_at = datetime.utcnow()
                    te.error_message = str(e)
                    self._add_trace(execution, "task.failed", agent_name=agent.name, task_name=task.name,
                                    data={"error": str(e)[:200]})
                    await db.commit()

                    # 发布任务失败事件
                    await event_publisher.publish_task_failed(
                        execution_id=self.execution_id,
                        task_name=task.name,
                        agent_name=agent.name,
                        error=str(e),
                    )

                    return False, local_total_tokens, local_total_cost

        return False, local_total_tokens, local_total_cost

    async def _execute_sub_crew(self, db, execution, task, te) -> tuple:
        """执行子工作流（嵌套 Crew）。

        创建子 Execution 并递归执行，将结果合并为当前任务的输出。
        """
        from app.models.crew import Crew
        from sqlalchemy.orm import selectinload

        self._add_trace(execution, "sub_crew.started", task_name=task.name, data={
            "sub_crew_id": task.sub_crew_id,
        })

        # 加载子 Crew
        result = await db.execute(
            select(Crew)
            .where(Crew.id == task.sub_crew_id)
            .options(
                selectinload(Crew.agents),
                selectinload(Crew.tasks),
            )
        )
        sub_crew = result.scalar_one_or_none()
        if not sub_crew:
            te.status = TaskExecutionStatus.FAILED
            te.error_message = f"子工作流 {task.sub_crew_id} 不存在"
            te.completed_at = datetime.utcnow()
            await db.commit()
            return False, 0, 0.0

        # 创建子 Execution
        sub_execution = Execution(
            crew_id=sub_crew.id,
            user_id=execution.user_id,
            trigger_type=execution.trigger_type,
            results={"parent_execution_id": execution.id, "parent_task_id": task.id},
        )
        db.add(sub_execution)
        await db.flush()
        await db.refresh(sub_execution)

        # 递归执行子 Crew
        sub_engine = ExecutionEngine(
            execution_id=sub_execution.id,
            llm_api_keys=self.llm_api_keys,
            llm_base_urls=self.llm_base_urls,
        )

        try:
            await sub_engine.run()

            # 刷新子执行状态
            await db.refresh(sub_execution)

            # 合并子执行结果
            sub_results = sub_execution.results or {}
            sub_outputs = sub_results.get("outputs", {})
            combined_output = "\n\n".join(
                f"[{tid}]: {output[:1000]}" for tid, output in sub_outputs.items()
            ) if sub_outputs else "(子工作流无输出)"

            te.status = TaskExecutionStatus.COMPLETED
            te.output = combined_output
            te.tokens_used = sub_execution.total_tokens_used or 0
            te.cost_usd = sub_execution.total_cost_usd or 0.0
            te.completed_at = datetime.utcnow()
            await db.commit()

            self._add_trace(execution, "sub_crew.completed", task_name=task.name, data={
                "sub_execution_id": sub_execution.id,
                "sub_status": sub_execution.status.value if hasattr(sub_execution.status, 'value') else str(sub_execution.status),
            })

            return True, te.tokens_used, te.cost_usd

        except Exception as e:
            te.status = TaskExecutionStatus.FAILED
            te.error_message = f"子工作流执行失败: {str(e)[:500]}"
            te.completed_at = datetime.utcnow()
            await db.commit()

            self._add_trace(execution, "sub_crew.failed", task_name=task.name, data={
                "error": str(e)[:200],
            })

            return False, 0, 0.0

    async def _handle_task_failure(self, db, execution, task, task_exec_map, error, total_tokens, total_cost, total_tasks):
        """处理任务失败"""
        execution.status = ExecutionStatus.FAILED
        execution.completed_at = datetime.utcnow()
        execution.error_log = f"任务'{task.name}'执行失败: {error}"
        execution.total_tokens_used = total_tokens
        execution.total_cost_usd = total_cost
        execution.results = {"completed_tasks": len([t for t in task_exec_map.values() if t.status == TaskExecutionStatus.COMPLETED]), "failed_task": task.name}
        await db.commit()

    async def _cancel_remaining(self, db, execution, task_exec_map):
        """取消剩余任务 — PENDING→SKIPPED, RUNNING→FAILED(已取消)"""
        execution.status = ExecutionStatus.CANCELLED
        execution.completed_at = datetime.utcnow()
        for te in task_exec_map.values():
            if te.status == TaskExecutionStatus.PENDING:
                te.status = TaskExecutionStatus.SKIPPED
            elif te.status == TaskExecutionStatus.RUNNING:
                te.status = TaskExecutionStatus.FAILED
                te.error_message = "任务已被用户取消"
                te.completed_at = datetime.utcnow()
        await db.commit()

    async def _fail_execution(self, db, execution, error):
        execution.status = ExecutionStatus.FAILED
        execution.completed_at = datetime.utcnow()
        execution.error_log = error
        await db.commit()

    @staticmethod
    def _describe_tool_call(tool_name: str, arguments: dict) -> str:
        """为工具调用生成自然语言推理描述"""
        desc_map = {
            "fs_read": lambda a: f"正在读取文件 '{a.get('path', '')}' 的内容...",
            "fs_write": lambda a: f"正在将结果写入文件 '{a.get('path', '')}'...",
            "fs_list": lambda a: f"正在查看 '{a.get('path', '.')}' 目录中的文件...",
            "docx_read": lambda a: f"正在读取 Word 文档 '{a.get('path', '')}'...",
            "docx_create": lambda a: f"正在创建 Word 文档 '{a.get('path', '')}'...",
            "xlsx_read": lambda a: f"正在读取 Excel 文件 '{a.get('path', '')}'...",
            "pdf_read": lambda a: f"正在读取 PDF 文件 '{a.get('path', '')}'...",
            "csv_analyze": lambda a: f"正在分析 CSV 文件 '{a.get('path', '')}'...",
            "file_search": lambda a: f"正在 '{a.get('directory', '.')}' 中搜索文件 '{a.get('pattern', '')}'...",
            "file_grep": lambda a: f"正在文件内容中搜索关键词 '{a.get('keyword', '')}'...",
            "web_search": lambda a: f"正在搜索互联网: '{a.get('query', '')}'...",
            "code_execute": lambda a: f"正在执行 {a.get('language', 'Python')} 代码...",
            "api_call": lambda a: f"正在调用 API: {a.get('method', 'GET')} {a.get('url', '')}...",
            "text_analysis": lambda a: f"正在分析文本 ({a.get('analysis_type', '')})...",
            "calculator": lambda a: f"正在计算: {a.get('expression', a.get('query', ''))}...",
            "date_time": lambda a: "正在获取当前日期和时间...",
            "hash_generate": lambda a: f"正在生成 {a.get('algorithm', 'sha256')} 哈希...",
            "process_info": lambda a: "正在获取系统信息...",
            "url_parse": lambda a: f"正在解析 URL: {a.get('url', '')[:60]}...",
            "json_query": lambda a: "正在查询 JSON 数据...",
            "yaml_parse": lambda a: f"正在解析 YAML 文件 '{a.get('path', '')}'...",
            "table_format": lambda a: "正在格式化表格数据...",
            "regex_extract": lambda a: f"正在用正则表达式提取数据: '{a.get('pattern', '')[:40]}'...",
            "text_diff": lambda a: "正在比较两段文本的差异...",
            "text_transform": lambda a: f"正在文本转换 ({a.get('operation', '')})...",
            "word_frequency": lambda a: "正在分析词频...",
            "zip_list": lambda a: f"正在查看 ZIP 文件 '{a.get('path', '')}' 的内容...",
            "sqlite_query": lambda a: "正在查询数据库...",
            "html_extract": lambda a: "正在从 HTML 中提取文本...",
            "markdown_render": lambda a: "正在将 Markdown 转换为 HTML...",
            "image_info": lambda a: f"正在获取图片信息 '{a.get('path', '')}'...",
            "image_generation": lambda a: f"正在生成图片: '{a.get('prompt', '')[:40]}'...",
            "csv_create": lambda a: f"正在创建 CSV 文件 '{a.get('path', '')}'...",
        }
        generator = desc_map.get(tool_name)
        if generator:
            try:
                return generator(arguments)
            except Exception as e:
                logger.warning(f"Tool call description failed for {tool_name}: {e}")
        return f"正在调用工具 {tool_name}..."

    @staticmethod
    def _parse_text_tool_calls(content: str) -> List['ToolCall']:
        """从 LLM 文本响应中解析 ````tool_call``` 代码块。

        用于不支持 function calling 的模型：模型在回复中嵌入
        ```tool_call
        {"tool": "docx_read", "args": {"path": "..."}}
        ```
        格式的调用指令，此方法将其解析为 ToolCall 对象。
        """
        from app.engine.llm_provider import ToolCall

        if not content:
            return []

        pattern = r'```tool_call\s*\n(.*?)\n```'
        matches = re.findall(pattern, content, re.DOTALL)

        tool_calls = []
        for match in matches:
            try:
                data = json.loads(match.strip())
                tool_name = data.get("tool", "")
                args = data.get("args", data.get("arguments", {}))
                if tool_name:
                    tool_calls.append(ToolCall(
                        id=f"text_tc_{uuid.uuid4().hex[:8]}",
                        name=tool_name,
                        arguments=args if isinstance(args, dict) else {},
                    ))
            except (json.JSONDecodeError, AttributeError):
                logger.warning(f"[EXECUTOR] Failed to parse text tool_call: {match[:200]}")
                continue

        return tool_calls

    async def _persist_memory_to_workspace(
        self, db, execution, crew, agents_map, task_outputs, total_tokens, total_cost
    ):
        """执行完成后，将结果记忆写入工作空间的 .fugue/ 目录"""
        import os as _os
        import traceback as _tb

        workspace = getattr(crew, 'workspace_dir', None)
        if not workspace:
            return {"status": "skipped", "reason": "workspace_dir 未设置"}

        ws_path = _os.path.abspath(_os.path.expanduser(workspace))
        af_dir = _os.path.join(ws_path, '.fugue')
        now = datetime.utcnow().isoformat()
        exec_id_short = execution.id[:8] if hasattr(execution, 'id') else 'unknown'
        created = []

        try:
            _os.makedirs(af_dir, exist_ok=True)

            # 1) 每个任务的完整输出保存为独立 .md 文件
            outputs_dir = _os.path.join(af_dir, 'outputs')
            _os.makedirs(outputs_dir, exist_ok=True)
            for task_id, output in task_outputs.items():
                out_file = _os.path.join(outputs_dir, f'{task_id[:8]}.md')
                with open(out_file, 'w', encoding='utf-8') as f:
                    f.write(str(output))
                created.append(f'outputs/{task_id[:8]}.md')

            # 2) 执行摘要
            lines = [f"# 执行记录 — {now}", "", f"- 执行ID: {exec_id_short}",
                     f"- 任务数: {len(task_outputs)}", f"- Token: {total_tokens}",
                     f"- 费用: ${total_cost:.4f}", "", "## 输出文件"]
            for tid in task_outputs:
                lines.append(f"- [outputs/{tid[:8]}.md](outputs/{tid[:8]}.md)")
            for a in (agents_map or {}).values():
                lines.append(f"- Agent **{a.name}**（{a.role}）")
            with open(_os.path.join(af_dir, 'last-execution.md'), 'w', encoding='utf-8') as f:
                f.write('\n'.join(lines))
            created.append('last-execution.md')

            # 3) Agent 经验
            for a in (agents_map or {}).values():
                af = _os.path.join(af_dir, f'agent-{a.name}.md')
                e = f"## {now}\n\n- 执行ID: {exec_id_short}\n- 角色: {a.role}\n- 目标: {a.goal}\n\n"
                ex = _os.path.exists(af)
                with open(af, 'a' if ex else 'w', encoding='utf-8') as f:
                    if not ex:
                        f.write(f'# Agent {a.name} — 经验记录\n\n')
                    f.write(e)
                created.append(f'agent-{a.name}.md')

            # 4) 项目总结
            sf = _os.path.join(af_dir, 'summary.md')
            if not _os.path.exists(sf):
                with open(sf, 'w', encoding='utf-8') as f:
                    f.write('# 项目执行总结\n\n| 时间 | 执行ID | 任务数 | Token | 费用 |\n|------|--------|--------|-------|------|\n')
            with open(sf, 'a', encoding='utf-8') as f:
                f.write(f"| {now} | {exec_id_short} | {len(task_outputs)} | {total_tokens} | ${total_cost:.4f} |\n")
            created.append('summary.md')

            # 5) AGENTS.md 放在 .fugue/ 内
            am = _os.path.join(af_dir, 'AGENTS.md')
            if not _os.path.exists(am):
                agents_content = (
                    f"# 项目约定 — {getattr(crew, 'name', '工作流')}\n\n"
                    "> 此文件在每次执行前自动注入到 Agent 上下文中。\n"
                    "> 编辑以添加项目编码规范、技术栈约定等。\n\n"
                    "## 技术栈\n\n待补充...\n\n## 编码规范\n\n待补充...\n\n## 注意事项\n\n待补充...\n"
                )
                with open(am, 'w', encoding='utf-8') as f:
                    f.write(agents_content)
                created.append('AGENTS.md')
                logger.info(f"[MEMORY] AGENTS.md created at {am}")
            else:
                logger.info(f"[MEMORY] AGENTS.md already exists, preserving user edits")

            logger.info(f"[MEMORY] {len(created)} files written to {af_dir}")
            return {"status": "ok", "path": af_dir, "files": created}

        except Exception as e:
            logger.error(f"[MEMORY] FAILED: {e}\n{_tb.format_exc()}")
            return {"status": "error", "reason": str(e)[:200]}

    async def _execute_event_flow(
        self, db, execution, crew, tasks, agents_map, task_map,
        task_outputs, task_exec_map, completed_task_ids,
        total_tokens, total_cost, completed_count, total_tasks,
    ):
        """Event Flow 模式执行 — 加载 FlowConfig，构建事件驱动流程并运行"""
        from app.models.flow_config import FlowConfig
        from app.engine.flow_executor import clear_flows, emit

        self._add_trace(execution, "crew.event_flow_start", data={"crew_id": crew.id})

        # 加载 FlowConfig
        result = await db.execute(
            select(FlowConfig).where(FlowConfig.crew_id == crew.id)
        )
        flow_configs = result.scalars().all()

        if not flow_configs:
            # 无 FlowConfig，回退到拓扑排序顺序执行
            self._add_trace(execution, "crew.event_flow_fallback", data={"reason": "no_flow_configs"})
            logger.warning(f"[EVENT_FLOW] No FlowConfigs for crew {crew.id}, falling back to sequential")
            levels = self._topological_levels(tasks, task_map)
            for level in levels:
                for task_id in level:
                    task = task_map.get(task_id)
                    if not task:
                        continue
                    success, t_tokens, t_cost = await self._execute_task(
                        db, execution, task, agents_map, task_outputs, task_exec_map
                    )
                    total_tokens += t_tokens
                    total_cost += t_cost
                    if success:
                        completed_count += 1
            return {
                "total_tokens": total_tokens,
                "total_cost": total_cost,
                "completed_count": completed_count,
                "flow_configs_count": 0,
            }

        # 按 flow_type 分组
        start_configs = [fc for fc in flow_configs if fc.flow_type == "start"]
        listen_configs = [fc for fc in flow_configs if fc.flow_type == "listen"]
        router_configs = [fc for fc in flow_configs if fc.flow_type == "router_event"]

        executed_task_ids = set()

        # Phase 1: 执行 @start 节点关联的任务
        for start_cfg in start_configs:
            linked_ids = start_cfg.linked_task_ids or []
            for task_id in linked_ids:
                task = task_map.get(task_id)
                if not task or task.id in executed_task_ids:
                    continue
                self._add_trace(execution, "event_flow.task_start", data={
                    "task_id": task_id, "trigger": start_cfg.event_name,
                })
                success, t_tokens, t_cost = await self._execute_task(
                    db, execution, task, agents_map, task_outputs, task_exec_map
                )
                total_tokens += t_tokens
                total_cost += t_cost
                if success:
                    completed_count += 1
                    executed_task_ids.add(task.id)
                # 触发后续事件
                if start_cfg.event_name and success:
                    await emit(start_cfg.event_name, data=task_outputs.get(task.id), source="event_flow")

        # Phase 2: 执行 @listen 节点关联的任务
        for listen_cfg in listen_configs:
            linked_ids = listen_cfg.linked_task_ids or []
            for task_id in linked_ids:
                task = task_map.get(task_id)
                if not task or task.id in executed_task_ids:
                    continue
                self._add_trace(execution, "event_flow.task_listen", data={
                    "task_id": task_id, "event": listen_cfg.event_name,
                })
                success, t_tokens, t_cost = await self._execute_task(
                    db, execution, task, agents_map, task_outputs, task_exec_map
                )
                total_tokens += t_tokens
                total_cost += t_cost
                if success:
                    completed_count += 1
                    executed_task_ids.add(task.id)

        # Phase 3: 执行 @router 节点关联的任务
        for router_cfg in router_configs:
            linked_ids = router_cfg.linked_task_ids or []
            for task_id in linked_ids:
                task = task_map.get(task_id)
                if not task or task.id in executed_task_ids:
                    continue
                self._add_trace(execution, "event_flow.task_router", data={
                    "task_id": task_id, "condition": router_cfg.condition,
                })
                success, t_tokens, t_cost = await self._execute_task(
                    db, execution, task, agents_map, task_outputs, task_exec_map
                )
                total_tokens += t_tokens
                total_cost += t_cost
                if success:
                    completed_count += 1
                    executed_task_ids.add(task.id)

        # Phase 4: 剩余未执行的任务（顺序兜底）
        for task in tasks:
            if task.id not in executed_task_ids:
                success, t_tokens, t_cost = await self._execute_task(
                    db, execution, task, agents_map, task_outputs, task_exec_map
                )
                total_tokens += t_tokens
                total_cost += t_cost
                if success:
                    completed_count += 1

        clear_flows()

        self._add_trace(execution, "crew.event_flow_complete", data={
            "flow_configs_count": len(flow_configs),
            "completed_count": completed_count,
        })
        return {
            "total_tokens": total_tokens,
            "total_cost": total_cost,
            "completed_count": completed_count,
            "flow_configs_count": len(flow_configs),
        }

    async def _execute_orchestrator_workers(
        self, db, execution, tasks, agents_map, task_map, task_outputs,
        task_exec_map, completed_task_ids, total_tokens, total_cost,
        completed_count, total_tasks, checkpoint_manager, loop_outputs,
    ) -> dict:
        """Orchestrator-Workers 模式：Orchestrator Agent 分解任务并动态分配给 Worker

        Task Ledger 记录每个子任务的分配、状态和结果。
        Worker 按角色匹配最适合的任务。
        """
        # 选取 Orchestrator（role 含 "orchestrator" 或 "manager" 的 Agent）
        orch_agent = None
        workers = []
        for a in agents_map.values():
            role_lower = (a.role or "").lower()
            if "orchestrator" in role_lower or "manager" in role_lower or "coordinator" in role_lower:
                orch_agent = a
            else:
                workers.append(a)
        if not orch_agent:
            orch_agent = workers[0] if workers else list(agents_map.values())[0]
            if orch_agent in workers:
                workers.remove(orch_agent)
        if not workers:
            workers = [orch_agent]  # fallback: 自己也是worker

        # Task Ledger: {task_id: {status, assigned_to, result}}
        ledger = {}
        for t in tasks:
            ledger[t.id] = {"task": t, "status": "pending", "assigned_to": None, "output": ""}

        self._add_trace(execution, "orchestrator.started", agent_name=orch_agent.name, data={
            "workers": [w.name for w in workers],
            "tasks": [t.name for t in tasks],
        })

        # 按任务角色匹配 Worker（简单关键词匹配）
        for task_id, entry in ledger.items():
            task = entry["task"]
            task_desc = (task.description or "").lower() + (task.name or "").lower()
            best_worker = workers[0]
            best_score = 0
            for w in workers:
                score = 0
                w_role = (w.role or "").lower()
                w_name = (w.name or "").lower()
                for kw in task_desc.split():
                    if kw in w_role or kw in w_name:
                        score += 1
                if score > best_score:
                    best_score = score
                    best_worker = w
            entry["assigned_to"] = best_worker.id
            entry["status"] = "assigned"

        # 按 Worker 分组执行任务
        worker_tasks = {}
        for task_id, entry in ledger.items():
            worker_id = entry["assigned_to"]
            worker_tasks.setdefault(worker_id, []).append(entry["task"])

        for worker_id, wtasks in worker_tasks.items():
            worker = agents_map.get(worker_id)
            if not worker:
                continue
            for task in wtasks:
                if self._cancelled:
                    break
                success, t_tokens, t_cost = await self._execute_task(
                    db, execution, task, agents_map, task_outputs, task_exec_map,
                )
                total_tokens += t_tokens
                total_cost += t_cost
                if success:
                    completed_count += 1
                    completed_task_ids.add(task.id)
                    ledger[task.id]["status"] = "completed"
                    ledger[task.id]["output"] = task_outputs.get(task.id, "")[:200]
                    await checkpoint_manager.create_auto_checkpoint(
                        task_id=task.id, task_name=task.name,
                        completed_task_ids=list(completed_task_ids),
                        task_outputs=dict(task_outputs),
                        total_tokens=total_tokens, total_cost_usd=total_cost,
                    )
                else:
                    ledger[task.id]["status"] = "failed"

        self._add_trace(execution, "orchestrator.completed", agent_name=orch_agent.name, data={
            "completed": completed_count,
            "total": total_tasks,
        })

        return {
            "total_tokens": total_tokens,
            "total_cost": total_cost,
            "completed_count": completed_count,
            "task_ledger": {k: {"status": v["status"], "worker": v["assigned_to"]} for k, v in ledger.items()},
        }

    # ── Agent 显式交接 (Command API) ──────────────────────────

    async def _handle_agent_command(
        self, db, execution, command_name: str, command_args: dict,
        current_agent, task, agents_map, task_outputs, task_exec_map,
    ) -> dict:
        """处理 Agent 间显式交接命令

        支持三种命令：
        - goto(target_agent) —— 将当前任务转交给目标 Agent
        - update(state) —— 更新共享状态
        - resume(value) —— 人机协作中注入反馈后恢复

        Returns:
            {handled: bool, message: str}
        """
        if command_name == "goto":
            target_name = command_args.get("target_agent", "")
            target = None
            for a in agents_map.values():
                if a.name == target_name or a.id == target_name:
                    target = a
                    break
            if not target:
                return {"handled": False, "message": f"Agent '{target_name}' not found"}

            # 将任务重新分配给目标 Agent
            old_agent_id = task.agent_id
            task.agent_id = target.id
            await db.flush()

            # 传递当前任务的输出作为新 Agent 的上下文
            handoff_context = command_args.get("context", "")
            if handoff_context:
                task_outputs[f"handoff_{current_agent.name}_to_{target.name}"] = handoff_context

            self._add_trace(execution, "agent.handoff", agent_name=current_agent.name, data={
                "from": current_agent.name,
                "to": target.name,
                "task": task.name,
            })
            logger.info(f"[HANDOFF] {current_agent.name} -> {target.name} for task '{task.name}'")
            return {"handled": True, "message": f"Task handed off to {target.name}"}

        elif command_name == "update":
            state_key = command_args.get("key", "")
            state_value = command_args.get("value", "")
            if state_key:
                execution.results = execution.results or {}
                if "shared_state" not in execution.results:
                    execution.results["shared_state"] = {}
                execution.results["shared_state"][state_key] = state_value
                await db.flush()
            return {"handled": True, "message": f"Updated shared state: {state_key}"}

        elif command_name == "resume":
            resume_value = command_args.get("value", "")
            task_outputs[f"resume_{task.id[:8]}"] = resume_value
            return {"handled": True, "message": "Resume value injected"}

        return {"handled": False, "message": f"Unknown command: {command_name}"}

    def _parse_agent_command(self, text: str) -> tuple:
        """从 LLM 输出中解析 Command API 调用

        格式: ```agent_command {"command": "goto", "target_agent": "..."}```
        """
        import re, json
        pattern = r'```agent_command\s*\n?(.*?)```'
        matches = re.findall(pattern, text, re.DOTALL)
        if not matches:
            return None, None
        try:
            cmd = json.loads(matches[0])
            return cmd.get("command"), cmd
        except (json.JSONDecodeError, KeyError):
            return None, None

    def _add_trace(self, execution, event_type, agent_name="", task_name="", data=None):
        if execution.trace is None:
            execution.trace = []
        execution.trace.append({
            "timestamp": datetime.utcnow().isoformat(),
            "event_type": event_type,
            "agent_name": agent_name,
            "task_name": task_name,
            "data": data or {},
        })
        flag_modified(execution, 'trace')

    async def _check_cost_alert(self, execution_id: str, total_cost: float):
        """检查成本是否接近预算上限，通过 WebSocket 发送预警"""
        budget = getattr(self, '_cost_budget', None)
        if not budget or budget <= 0:
            return

        ratio = total_cost / budget
        level = None
        if ratio >= 1.0:
            level = "exceeded"
        elif ratio >= 0.9:
            level = "critical"
        elif ratio >= 0.8:
            level = "warning"

        if level:
            await event_publisher.publish_warning(
                execution_id=execution_id,
                message=f"成本已达预算的 {ratio:.0%}（${total_cost:.4f} / ${budget:.2f}）",
                level=level,
            )
            logger.warning(
                f"[EXECUTOR] Cost alert: {level} - ${total_cost:.4f}/{budget:.2f} ({ratio:.0%}) "
                f"for execution {execution_id}"
            )

    async def _build_memory_context(self, db, agent, task, execution) -> str:
        """构建记忆上下文（委托给 prompt_builder）"""
        from app.engine.prompt_builder import build_memory_context
        memory_config = getattr(self, '_memory_config', None)
        return await build_memory_context(db, agent, task, execution, memory_config)

    def _build_messages(self, agent, task, context_parts):
        """构建 LLM 消息（委托给 prompt_builder）"""
        from app.engine.prompt_builder import build_messages
        workspace = getattr(self, '_workspace_dir', None)
        return build_messages(agent, task, context_parts, workspace_dir=workspace)

    def _topological_levels(self, tasks, task_map):
        """将任务按拓扑层级分组（同层任务可并行执行）"""
        in_degree: Dict[str, int] = defaultdict(int)
        dependents: Dict[str, List[str]] = defaultdict(list)

        for task in tasks:
            deps = [d for d in (task.context_task_ids or []) if d in task_map]
            in_degree[task.id] = len(deps)
            for dep_id in deps:
                dependents[dep_id].append(task.id)

        # BFS分层
        levels = []
        current_level = [t for t in tasks if in_degree[t.id] == 0]

        while current_level:
            levels.append(current_level)
            next_level = []
            for task in current_level:
                for dependent_id in dependents[task.id]:
                    in_degree[dependent_id] -= 1
                    if in_degree[dependent_id] == 0:
                        next_level.append(task_map[dependent_id])
            current_level = next_level

        # 检测环（如果排序后的任务数 < 总任务数，说明有环）
        sorted_count = sum(len(level) for level in levels)
        if sorted_count < len(tasks):
            remaining = [t for t in tasks if t.id not in {t2.id for level in levels for t2 in level}]
            logger.warning(f"检测到循环依赖，剩余任务按创建时间排序: {[t.name for t in remaining]}")
            levels.append(sorted(remaining, key=lambda t: t.created_at or datetime.min))

        return levels

    def _topological_sort(self, tasks):
        """向后兼容：返回扁平排序列表"""
        task_map = {t.id: t for t in tasks}
        levels = self._topological_levels(tasks, task_map)
        return [task for level in levels for task in level]

    async def _execute_condition(
        self,
        condition: ConditionBranch,
        context: Dict[str, Any],
    ) -> List[str]:
        """执行条件分支（委托给 conditions 模块）"""
        from app.engine.conditions import evaluate_condition
        return evaluate_condition(condition, context)

    async def execute_with_conditions(
        self,
        tasks: List[Task],
        conditions: List[ConditionBranch],
        context: Dict[str, Any],
    ) -> Dict[str, Any]:
        """执行支持条件分支的工作流

        Args:
            tasks: 任务列表
            conditions: 条件分支列表
            context: 执行上下文

        Returns:
            执行结果字典
        """
        task_map = {t.id: t for t in tasks}
        condition_map = {c.id: c for c in conditions}

        # 构建执行图：根据条件分支决定要执行的任务
        executable_task_ids: Set[str] = set()

        for condition in conditions:
            branch_task_ids = await self._execute_condition(condition, context)
            for tid in branch_task_ids:
                if tid in task_map:
                    executable_task_ids.add(tid)

        # 无条件关联的任务默认执行
        conditioned_task_ids: Set[str] = set()
        for condition in conditions:
            conditioned_task_ids.update(condition.true_branch_task_ids or [])
            conditioned_task_ids.update(condition.false_branch_task_ids or [])

        for task in tasks:
            if task.id not in conditioned_task_ids:
                executable_task_ids.add(task.id)

        executable_tasks = [task_map[tid] for tid in executable_task_ids if tid in task_map]
        return {"executable_tasks": executable_tasks, "skipped_tasks": [t for t in tasks if t.id not in executable_task_ids]}

    async def _execute_loop(
        self,
        loop_config: LoopConfig,
        tasks: List[Task],
        agents_map: Dict[str, Agent],
        execution: Execution,
        task_outputs: Dict[str, str],
        task_exec_map: Dict[str, TaskExecution],
        db: AsyncSession,
        checkpoint_manager: 'CheckpointManager',
        total_tokens: int,
        total_cost: float,
        completed_task_ids: Set[str],
        completed_count: int,
        total_tasks: int,
    ) -> Dict[str, Any]:
        """执行循环

        Args:
            loop_config: 循环配置
            tasks: 全部任务列表
            agents_map: Agent映射
            execution: 执行记录
            task_outputs: 任务输出映射
            task_exec_map: 任务执行映射
            db: 数据库会话
            checkpoint_manager: 断点管理器
            total_tokens: 已累计token
            total_cost: 已累计费用
            completed_task_ids: 已完成任务ID集合
            completed_count: 已完成任务计数
            total_tasks: 总任务数

        Returns:
            循环执行结果
        """
        task_map = {t.id: t for t in tasks}
        loop_tasks = [task_map[tid] for tid in (loop_config.loop_body_task_ids or []) if tid in task_map]

        results = {}
        iteration = 0

        self._add_trace(execution, "loop.started", task_name=loop_config.name, data={
            "max_iterations": loop_config.max_iterations,
            "tasks": [t.name for t in loop_tasks],
        })

        while iteration < loop_config.max_iterations:
            # 检查是否取消
            if self._cancelled:
                self._add_trace(execution, "loop.cancelled", task_name=loop_config.name, data={"iteration": iteration})
                break

            iteration += 1

            # 检查是否应该继续循环
            if loop_config.condition and iteration > 1:
                try:
                    safe_names = {
                        "context": task_outputs,
                        "iteration": iteration,
                        "results": results,
                    }
                    safe_functions = {
                        "len": len,
                        "str": str,
                        "int": int,
                        "float": float,
                        "bool": bool,
                    }
                    should_continue = simple_eval(loop_config.condition, names=safe_names, functions=safe_functions)
                    if not should_continue:
                        self._add_trace(execution, "loop.condition_false", task_name=loop_config.name, data={
                            "iteration": iteration,
                            "condition": loop_config.condition,
                        })
                        break
                except InvalidExpression as e:
                    if loop_config.exit_on_failure:
                        raise ValueError(f"循环条件无效: {str(e)}")
                    self._add_trace(execution, "loop.condition_error", task_name=loop_config.name, data={
                        "iteration": iteration,
                        "error": str(e),
                    })
                    break
                except Exception as e:
                    if loop_config.exit_on_failure:
                        raise ValueError(f"循环条件评估失败: {str(e)}")
                    self._add_trace(execution, "loop.condition_error", task_name=loop_config.name, data={
                        "iteration": iteration,
                        "error": str(e),
                    })
                    break

            # 执行循环体任务
            self._add_trace(execution, "loop.iteration_start", task_name=loop_config.name, data={
                "iteration": iteration,
            })

            try:
                is_parallel = False
                # 按拓扑排序执行循环体任务
                levels = self._topological_levels(loop_tasks, task_map)

                for level_tasks in levels:
                    if self._cancelled:
                        break

                    pending_tasks = [
                        t for t in level_tasks
                        if t.id not in completed_task_ids or iteration > 1
                    ]

                    for task in pending_tasks:
                        if self._cancelled:
                            break

                        success, task_tokens, task_cost = await self._execute_task(
                            db, execution, task, agents_map, task_outputs, task_exec_map,
                        )
                        if not success:
                            if loop_config.exit_on_failure:
                                raise RuntimeError(f"循环体任务'{task.name}'执行失败")
                            results[f"iteration_{iteration}"] = {"error": f"任务'{task.name}'失败"}
                            break
                        total_tokens += task_tokens
                        total_cost += task_cost
                        if task.id not in completed_task_ids:
                            completed_count += 1
                            completed_task_ids.add(task.id)

                iteration_result = {
                    "outputs": {t.name: task_outputs.get(t.id, "") for t in loop_tasks},
                }
                results[f"iteration_{iteration}"] = iteration_result

                # 更新上下文供下次迭代使用
                self._add_trace(execution, "loop.iteration_complete", task_name=loop_config.name, data={
                    "iteration": iteration,
                })

            except Exception as e:
                if loop_config.exit_on_failure:
                    raise
                results[f"iteration_{iteration}"] = {"error": str(e)}
                self._add_trace(execution, "loop.iteration_error", task_name=loop_config.name, data={
                    "iteration": iteration,
                    "error": str(e),
                })

        self._add_trace(execution, "loop.completed", task_name=loop_config.name, data={
            "iterations": iteration,
        })

        return {
            "iterations": iteration,
            "results": results,
            "loop_config_id": loop_config.id,
            "loop_name": loop_config.name,
            "total_tokens": total_tokens,
            "total_cost": total_cost,
            "completed_count": completed_count,
        }

    async def execute_with_loops(
        self,
        tasks: List[Task],
        loop_configs: List[LoopConfig],
        agents_map: Dict[str, Agent],
        execution: Execution,
        task_outputs: Dict[str, str],
        task_exec_map: Dict[str, TaskExecution],
        db: AsyncSession,
        checkpoint_manager: 'CheckpointManager',
        total_tokens: int,
        total_cost: float,
        completed_task_ids: Set[str],
        completed_count: int,
        total_tasks: int,
    ) -> Dict[str, Any]:
        """执行支持循环的工作流

        Args:
            tasks: 任务列表
            loop_configs: 循环配置列表
            agents_map: Agent映射
            execution: 执行记录
            task_outputs: 任务输出映射
            task_exec_map: 任务执行映射
            db: 数据库会话
            checkpoint_manager: 断点管理器
            total_tokens: 已累计token
            total_cost: 已累计费用
            completed_task_ids: 已完成任务ID集合
            completed_count: 已完成任务计数
            total_tasks: 总任务数

        Returns:
            包含循环结果和更新后的执行状态
        """
        task_map = {t.id: t for t in tasks}
        loop_task_ids: Set[str] = set()

        # 收集所有循环体任务ID
        for lc in loop_configs:
            loop_task_ids.update(lc.loop_body_task_ids or [])

        # 分离循环任务和普通任务
        loop_results = []
        for loop_config in loop_configs:
            result = await self._execute_loop(
                loop_config=loop_config,
                tasks=tasks,
                agents_map=agents_map,
                execution=execution,
                task_outputs=task_outputs,
                task_exec_map=task_exec_map,
                db=db,
                checkpoint_manager=checkpoint_manager,
                total_tokens=total_tokens,
                total_cost=total_cost,
                completed_task_ids=completed_task_ids,
                completed_count=completed_count,
                total_tasks=total_tasks,
            )
            loop_results.append(result)

            # 更新累计值（从循环结果中获取更新后的值）
            total_tokens = result.get("total_tokens", total_tokens)
            total_cost = result.get("total_cost", total_cost)
            completed_count = result.get("completed_count", completed_count)

        # 返回非循环任务（由调用方继续执行）
        non_loop_tasks = [t for t in tasks if t.id not in loop_task_ids]

        return {
            "loop_results": loop_results,
            "non_loop_tasks": non_loop_tasks,
            "total_tokens": total_tokens,
            "total_cost": total_cost,
            "completed_count": completed_count,
            "completed_task_ids": completed_task_ids,
        }

    async def run_iteration(self, execution_id: str, iteration_id: str):
        """执行迭代优化"""
        from app.models.iteration import Iteration, IterationStatus

        async with get_db_session() as db:
            iteration = await db.get(Iteration, iteration_id)
            if not iteration:
                logger.error(f"Iteration {iteration_id} not found")
                raise ValueError(f"Iteration {iteration_id} not found")

            # 从关联的 agent 获取模型名称 + 工作空间 + 工具配置
            execution = await db.get(Execution, execution_id)
            agent_tool_names = []
            if execution:
                result = await db.execute(
                    select(Crew).where(Crew.id == execution.crew_id)
                    .options(selectinload(Crew.agents))
                )
                crew = result.scalar_one_or_none()
                if crew and crew.agents:
                    self._iteration_model = crew.agents[0].llm_model
                    agent_tool_names = crew.agents[0].tools_config or []
                else:
                    self._iteration_model = None
                # 设置工作空间（迭代也要遵守工作空间限制）
                if crew and crew.workspace_dir:
                    self._workspace_dir = crew.workspace_dir
                    from app.engine.tools import set_workspace_dir
                    set_workspace_dir(crew.workspace_dir)
                    logger.info(f"[ITERATION] Workspace set to: {crew.workspace_dir}")
            else:
                self._iteration_model = None

            iteration.status = IterationStatus.RUNNING
            await db.commit()

            # 安全获取 mode 值（兼容 enum 和 string）
            mode_val = iteration.mode.value if hasattr(iteration.mode, 'value') else str(iteration.mode)

            try:
                if mode_val == "reexecute":
                    result = await self._reexecute_with_feedback(
                        execution_id, iteration.feedback, agent_tool_names=agent_tool_names
                    )
                else:
                    result = await self._incremental_refine(
                        iteration.previous_output or "", iteration.feedback,
                        execution_id=execution_id, agent_tool_names=agent_tool_names,
                    )

                iteration.refined_output = result['output']
                iteration.tokens_used = result['tokens_used']
                iteration.cost_usd = result['cost_usd']
                iteration.status = IterationStatus.COMPLETED
                iteration.completed_at = datetime.utcnow()
                logger.info(f"Iteration {iteration_id} completed: {len(result['output'])} chars, {result['tokens_used']} tokens")

            except Exception as e:
                logger.error(f"Iteration {iteration_id} failed: {e}", exc_info=True)
                iteration.status = IterationStatus.FAILED
                iteration.error_message = str(e)[:500]
            finally:
                # 清理工作空间引用，防止影响后续执行
                from app.engine.tools import clear_workspace_dir
                clear_workspace_dir()

            await db.commit()

    async def _incremental_refine(self, previous_output: str, feedback: str, execution_id: str = None, agent_tool_names: list = None) -> dict:
        """增量优化：在上次结果上修改（支持工具调用）

        查询当前 execution 的所有历史迭代（按 created_at 排序），
        构建累积上下文：将每次迭代的 feedback + refined_output 按时间顺序拼接。
        """
        # 构建累积历史上下文
        history_context = ""
        if execution_id:
            try:
                async with get_db_session() as db:
                    from app.models.iteration import Iteration
                    history_result = await db.execute(
                        select(Iteration)
                        .where(Iteration.execution_id == execution_id)
                        .order_by(Iteration.created_at.asc())
                    )
                    historical_iterations = history_result.scalars().all()

                    if historical_iterations:
                        history_parts = []
                        for idx, it in enumerate(historical_iterations, 1):
                            mode_str = it.mode.value if hasattr(it.mode, 'value') else str(it.mode)
                            part = f"""--- 迭代 #{idx} (模式: {mode_str}) ---
反馈: {it.feedback}
结果: {it.refined_output or '(无输出)'}"""
                            history_parts.append(part)
                        history_context = "\n\n".join(history_parts)
            except Exception as e:
                logger.warning(f"[ITERATION] Failed to load iteration history: {e}")

        prompt_parts = []

        # 如果有累积历史，注入完整历史链
        if history_context:
            prompt_parts.append(f"""历史迭代记录（按时间顺序）：
{history_context}""")

        # 当前迭代内容
        prompt_parts.append(f"""当前迭代的上次输出：
{previous_output}

用户反馈：
{feedback}""")

        if history_context:
            prompt_parts.append(
                "请综合以上所有历史迭代记录和最新反馈，对输出进行优化。"
                "注意避免重复之前已经尝试过的修改方向。"
                "如果需要读取文件或保存结果，请使用工具。"
            )
        else:
            prompt_parts.append(
                "请根据反馈修改输出。如果需要读取文件或保存结果，请使用工具。"
            )

        prompt = "\n\n".join(prompt_parts)

        return await self._call_llm_with_tools(prompt, execution_id=execution_id, agent_tool_names=agent_tool_names)

    async def _reexecute_with_feedback(self, execution_id: str, feedback: str, agent_tool_names: list = None) -> dict:
        """重新执行模式：提取原始任务描述并附加反馈重新执行（支持工具调用）"""
        async with get_db_session() as db:
            execution = await db.get(Execution, execution_id)
            if not execution:
                raise ValueError(f"执行 {execution_id} 不存在")

            # 获取关联的 crew 和任务
            result = await db.execute(
                select(Crew)
                .where(Crew.id == execution.crew_id)
                .options(selectinload(Crew.tasks), selectinload(Crew.agents))
            )
            crew = result.scalar_one_or_none()
            if not crew:
                raise ValueError("关联工作流不存在")

            # 收集所有任务的描述作为上下文
            task_summaries = []
            for task in (crew.tasks or []):
                task_summaries.append(f"- {task.name}: {task.description or '无描述'}")

            tasks_text = "\n".join(task_summaries) if task_summaries else "无任务"

            # 从 task_executions 构建完整上下文
            te_result = await db.execute(
                select(TaskExecution)
                .where(TaskExecution.execution_id == execution_id)
                .order_by(TaskExecution.created_at.asc())
            )
            task_executions = te_result.scalars().all()

            # 构建任务名称映射
            task_map = {t.id: t.name for t in (crew.tasks or [])}

            # 优先使用 task_executions 构建完整上下文
            if task_executions:
                output_parts = []
                for te in task_executions:
                    task_name = task_map.get(te.task_id, te.task_id)
                    te_status = te.status.value if hasattr(te.status, 'value') else str(te.status)
                    output_parts.append(
                        f"[{task_name}] (状态: {te_status}, tokens: {te.tokens_used or 0}):\n"
                        f"{(te.output or '(无输出)')[:3000]}"
                    )
                previous_output = "\n\n".join(output_parts)
            else:
                # 回退到 execution.results
                outputs = (execution.results or {}).get("outputs", {})
                previous_output = "\n\n".join(
                    f"[{tid}]: {out[:2000]}" for tid, out in outputs.items()
                ) if outputs else "无上次执行结果"

            prompt = f"""请根据用户反馈重新生成工作流的执行结果。

工作流任务：
{tasks_text}

上次执行完整结果：
{previous_output}

用户反馈：
{feedback}

请根据反馈生成改进后的完整结果。如果需要读取文件或保存结果，请使用工具。"""

        return await self._call_llm_with_tools(prompt, execution_id=execution_id, agent_tool_names=agent_tool_names)

    async def _get_agent_context(self, execution_id: str = None) -> str:
        """查询 execution 关联的 crew 中第一个 agent，构建角色上下文注入 system prompt。

        Returns:
            agent 角色信息字符串，无 agent 时返回默认前缀。
        """
        if not execution_id:
            return "你是一个专业的AI助手。"

        try:
            async with get_db_session() as db:
                execution = await db.get(Execution, execution_id)
                if not execution:
                    return "你是一个专业的AI助手。"

                result = await db.execute(
                    select(Crew)
                    .where(Crew.id == execution.crew_id)
                    .options(selectinload(Crew.agents))
                )
                crew = result.scalar_one_or_none()
                if not crew or not crew.agents:
                    return "你是一个专业的AI助手。"

                agent = crew.agents[0]
                parts = [f"你正在扮演 {agent.name}"]
                if agent.role:
                    parts.append(f"角色：{agent.role}")
                if agent.backstory:
                    parts.append(f"背景：{agent.backstory}")
                if agent.goal:
                    parts.append(f"目标：{agent.goal}")
                return "。".join(parts) + "。\n"
        except Exception as e:
            logger.warning(f"[ITERATION] Failed to load agent context: {e}")
            return "你是一个专业的AI助手。"

    async def _call_llm_with_tools(self, prompt: str, model: str = None, execution_id: str = None, agent_tool_names: list = None) -> dict:
        """调用 LLM 并支持工具调用循环（迭代优化专用）

        与 _call_llm 不同：此方法传入工具 schema，支持多轮工具调用。
        当提供 execution_id 时，会查询关联 agent 角色信息注入 system prompt。
        """
        from app.engine.tools import get_plugin_tool_schemas, execute_tool, get_openai_tools, get_anthropic_tools

        # 获取 LLM provider
        all_keys = self.llm_api_keys or {}
        all_urls = self.llm_base_urls or {}
        provider_name = None
        api_key = None
        base_url = None
        for pname, pkey in all_keys.items():
            if pkey:
                provider_name = pname
                api_key = pkey
                base_url = all_urls.get(pname)
                break

        if not provider_name:
            from app.core.config import settings
            if settings.OPENAI_API_KEY:
                provider_name, api_key = "openai", settings.OPENAI_API_KEY
            elif settings.ANTHROPIC_API_KEY:
                provider_name, api_key = "anthropic", settings.ANTHROPIC_API_KEY

        if not provider_name or not api_key:
            return {"output": f"[演示模式]\n{prompt[:500]}", "tokens_used": 0, "cost_usd": 0.0}

        llm = get_llm_provider(provider_name, api_key=api_key, base_url=base_url, all_keys=all_keys, all_base_urls=all_urls)
        from app.engine.llm_provider import MockProvider
        if isinstance(llm, MockProvider):
            return {"output": f"[演示模式]\n{prompt[:500]}", "tokens_used": 0, "cost_usd": 0.0}

        if not model:
            model = self._iteration_model or "gpt-4o"

        # 构建工具 schema — 与 _execute_task 保持一致，加载全部工具源
        is_anthropic = provider_name == "anthropic"
        provider_fmt = "anthropic" if is_anthropic else "openai"
        tool_schemas = []

        try:
            # 1) 内置工具（按 agent 配置筛选）
            if agent_tool_names:
                tool_schemas = get_anthropic_tools(agent_tool_names) if is_anthropic else get_openai_tools(agent_tool_names)

            # 2) MCP 工具
            try:
                from app.engine.mcp_adapter import get_mcp_adapter
                mcp_schemas = get_mcp_adapter().get_tool_schemas()
                if mcp_schemas:
                    tool_schemas = (tool_schemas or []) + mcp_schemas
            except Exception as e:
                logger.debug(f"[ITERATION] MCP tools unavailable: {e}")

            # 3) 插件工具（LocalFS + EnhancedTools 等）
            plugin_schemas = get_plugin_tool_schemas(provider_fmt)
            if plugin_schemas:
                tool_schemas = (tool_schemas or []) + plugin_schemas

            # 4) 兜底：如果无任何工具，提供默认内置工具集
            if not agent_tool_names and not tool_schemas:
                _default_tools = ["web_search", "file_read", "file_write", "code_execute", "api_call", "text_analysis"]
                builtin_schemas = get_anthropic_tools(_default_tools) if is_anthropic else get_openai_tools(_default_tools)
                if builtin_schemas:
                    tool_schemas = builtin_schemas

            logger.info(
                f"[ITERATION] Tool schemas: "
                f"agent_config={len(agent_tool_names or [])}, "
                f"plugin={len(plugin_schemas) if plugin_schemas else 0}, "
                f"total={len(tool_schemas or [])}"
            )
        except Exception as e:
            logger.warning(f"[ITERATION] Failed to load tool schemas, continuing without tools: {e}")
            tool_schemas = []

        # 构建 system prompt（注入 agent 角色信息）
        agent_context = await self._get_agent_context(execution_id)
        system_content = agent_context + (
            "根据反馈优化输出。如果需要读取文件或保存结果到文件，请主动使用工具。保存文件请使用 fs_write。"
        )

        messages = [
            {"role": "system", "content": system_content},
            {"role": "user", "content": prompt},
        ]

        total_tokens = 0
        total_cost = 0.0
        final_content = ""

        # 工具调用循环（最多 5 轮）
        tools_failed = False  # 标记工具是否被 API 拒绝
        _LLM_TIMEOUT = 120  # 每次 LLM 调用超时 120 秒
        for tool_round in range(5):
            try:
                use_tools = tool_schemas if (tool_schemas and not tools_failed) else None
                resp = await asyncio.wait_for(
                    llm.chat(
                        messages=messages, model=model, temperature=0.7, max_tokens=4096,
                        tools=use_tools,
                    ),
                    timeout=_LLM_TIMEOUT,
                )
            except asyncio.TimeoutError:
                logger.error(f"[ITERATION] LLM call timed out after {_LLM_TIMEOUT}s (round {tool_round+1})")
                raise RuntimeError(f"LLM 调用超时（{_LLM_TIMEOUT}秒），请检查网络或 API 状态")
            except Exception as e:
                # 如果带 tools 的请求失败（400），重试不带 tools
                if tool_schemas and not tools_failed and "400" in str(e):
                    logger.warning(f"[ITERATION] Tools rejected by API, retrying without tools: {e}")
                    tools_failed = True
                    try:
                        resp = await asyncio.wait_for(
                            llm.chat(
                                messages=messages, model=model, temperature=0.7, max_tokens=4096,
                                tools=None,
                            ),
                            timeout=_LLM_TIMEOUT,
                        )
                    except asyncio.TimeoutError:
                        logger.error(f"[ITERATION] LLM call timed out after {_LLM_TIMEOUT}s (no-tools retry)")
                        raise RuntimeError(f"LLM 调用超时（{_LLM_TIMEOUT}秒）")
                    except Exception as e2:
                        logger.error(f"[ITERATION] LLM call failed even without tools: {e2}")
                        raise
                else:
                    logger.error(f"[ITERATION] LLM call failed: {e}")
                    raise

            total_tokens += resp.tokens_used
            total_cost += resp.cost_usd

            if not resp.tool_calls:
                final_content = resp.content or ""
                break

            # 执行工具（防御性处理，单个工具失败不影响其他）
            for tc in resp.tool_calls:
                try:
                    tool_result = await execute_tool(
                        tc.name or "unknown",
                        tc.arguments if isinstance(tc.arguments, dict) else {},
                        tc.id or "",
                    )
                    tc.result = tool_result.output
                except Exception as tool_err:
                    logger.error(f"[ITERATION] Tool {tc.name} failed: {tool_err}")
                    tc.result = f"[错误] 工具执行失败: {tool_err}"

            # 追加消息历史
            if is_anthropic:
                assistant_blocks = []
                if resp.content:
                    assistant_blocks.append({"type": "text", "text": resp.content})
                for tc in resp.tool_calls:
                    assistant_blocks.append({"type": "tool_use", "id": tc.id, "name": tc.name, "input": tc.arguments})
                messages.append({"role": "assistant", "content": assistant_blocks})
                tool_result_blocks = []
                for tc in resp.tool_calls:
                    tool_result_blocks.append({"type": "tool_result", "tool_use_id": tc.id, "content": (tc.result or "")[:8000]})
                messages.append({"role": "user", "content": tool_result_blocks})
            else:
                assistant_msg = {"role": "assistant", "content": resp.content or ""}
                assistant_msg["tool_calls"] = [
                    {"id": tc.id, "type": "function", "function": {"name": tc.name, "arguments": json.dumps(tc.arguments, ensure_ascii=False)}}
                    for tc in resp.tool_calls
                ]
                messages.append(assistant_msg)
                for tc in resp.tool_calls:
                    messages.append({"role": "tool", "tool_call_id": tc.id, "content": (tc.result or "")[:8000]})
        else:
            # 循环用完，强制请求最终回复
            try:
                messages.append({"role": "user", "content": "请基于以上工具调用结果，给出最终答案。不要再调用工具。"})
                resp = await llm.chat(messages=messages, model=model, temperature=0.7, max_tokens=4096, tools=None)
                final_content = resp.content or ""
                total_tokens += resp.tokens_used
                total_cost += resp.cost_usd
            except Exception:
                final_content = ""

        if not final_content:
            final_content = "(迭代优化完成，但未生成最终内容)"

        return {"output": final_content, "tokens_used": total_tokens, "cost_usd": total_cost}

    async def _call_llm(self, prompt: str, model: str = None) -> dict:
        """调用 LLM 进行迭代优化（使用执行引擎已配置的 API keys）

        Args:
            prompt: 用户提示
            model: 模型名称（如不传则从执行上下文获取）
        """
        # 使用引擎已存储的 LLM 配置
        all_keys = self.llm_api_keys or {}
        all_urls = self.llm_base_urls or {}

        # 选择第一个可用的 provider
        provider_name = None
        api_key = None
        base_url = None
        for pname, pkey in all_keys.items():
            if pkey:
                provider_name = pname
                api_key = pkey
                base_url = all_urls.get(pname)
                break

        # 如果没有用户传入的 key，尝试从环境变量获取
        if not provider_name:
            from app.core.config import settings
            if settings.OPENAI_API_KEY:
                provider_name = "openai"
                api_key = settings.OPENAI_API_KEY
            elif settings.ANTHROPIC_API_KEY:
                provider_name = "anthropic"
                api_key = settings.ANTHROPIC_API_KEY

        if not provider_name or not api_key:
            return {
                "output": f"[迭代优化 - 演示模式]\n\n原始内容：\n{prompt[:500]}\n\n提示：请配置 LLM API Key 以获得真实的迭代优化结果。",
                "tokens_used": 0,
                "cost_usd": 0.0,
            }

        llm = get_llm_provider(
            provider_name,
            api_key=api_key,
            base_url=base_url,
            all_keys=all_keys,
            all_base_urls=all_urls,
        )

        messages = [
            {"role": "system", "content": "你是一个专业的AI助手，负责根据用户反馈优化输出内容。请直接返回优化后的内容，不要添加额外解释。"},
            {"role": "user", "content": prompt},
        ]

        from app.engine.llm_provider import MockProvider
        if isinstance(llm, MockProvider):
            return {
                "output": f"[迭代优化 - 演示模式]\n\n原始内容：\n{prompt[:500]}\n\n提示：请配置有效的 LLM API Key。",
                "tokens_used": 0,
                "cost_usd": 0.0,
            }

        # 使用传入的模型名，或从 agent 配置获取，或用 provider 默认模型
        if not model:
            model = self._iteration_model or "gpt-4o"

        try:
            response = await llm.chat(
                messages=messages,
                model=model,
                temperature=0.7,
                max_tokens=4096,
            )
            return {
                "output": response.content or "(LLM 未返回内容)",
                "tokens_used": response.tokens_used,
                "cost_usd": response.cost_usd,
            }
        except Exception as e:
            logger.error(f"[ITERATION] LLM call failed with model '{model}': {e}")
            raise


# 全局执行引擎注册表
_running_engines: Dict[str, ExecutionEngine] = {}


async def start_execution(
    execution_id: str,
    llm_api_keys: Dict[str, str] = None,
    llm_base_urls: Dict[str, str] = None,
    resume: bool = False,
    max_turns: int = 10,
):
    logger.info(f"[START_EXECUTION] Starting execution engine for {execution_id}, resume={resume}, max_turns={max_turns}")
    engine = ExecutionEngine(execution_id, llm_api_keys, llm_base_urls, max_turns=max_turns)
    _running_engines[execution_id] = engine
    try:
        logger.info(f"[START_EXECUTION] Calling engine.run(resume={resume}) for {execution_id}")
        await engine.run(resume=resume)
        logger.info(f"[START_EXECUTION] engine.run() completed for {execution_id}")
    except Exception as e:
        logger.error(f"[START_EXECUTION] Execution {execution_id[:8]} error: {e}", exc_info=True)
        raise  # 重新抛出，让调用方（_run_execution_background）能捕获并标记 FAILED
    finally:
        _running_engines.pop(execution_id, None)
        logger.info(f"[START_EXECUTION] Execution {execution_id} removed from running engines")


    # ═══════════════════════════════════════════════════════════════
    # 新增 ProcessType 实现（报告第1.1, 1.4节）
    # ═══════════════════════════════════════════════════════════════

    async def _execute_plan_execute(
        self, db, execution, tasks, agents_map, task_map,
        task_outputs, task_exec_map, completed_task_ids,
        total_tokens, total_cost, completed_count, total_tasks,
    ) -> Dict[str, Any]:
        """Plan-Execute 模式：Agent 先规划再执行。

        1. Plan Phase — 取第一个 Agent 作为 Planner，生成结构化计划
        2. Execute Phase — 按计划逐步调用工具执行
        """
        logger.info(f"[PLAN_EXECUTE] Starting with {len(tasks)} task(s), {len(agents_map)} agent(s)")

        for task in sorted(tasks, key=lambda t: t.order if hasattr(t, 'order') else 0):
            if self._cancelled or self._paused:
                break
            if task.id in completed_task_ids:
                continue

            agent = agents_map.get(task.agent_id)
            if not agent:
                continue

            # Plan Phase: 让 LLM 生成执行计划
            plan_prompt = (
                "在开始执行之前，请先制定一个详细的执行计划。\n"
                "以 JSON 格式输出，包含以下字段：\n"
                '{"steps": [{"step": 1, "action": "描述要做什么", "tool": "工具名（可选）", "expected": "预期结果"}]}\n'
                "然后逐步执行你的计划。"
            )
            task_with_plan = type('_PlanTask', (), {
                'id': task.id, 'name': task.name,
                'description': f"{plan_prompt}\n\n原始任务：{task.description}",
                'expected_output': task.expected_output, 'agent_id': task.agent_id,
                'timeout_seconds': getattr(task, 'timeout_seconds', 300),
                'max_retries': getattr(task, 'max_retries', 1),
            })()

            te = task_exec_map.get(task.id)
            result = await self._execute_task(
                db, execution, task_with_plan, agent, te, task_outputs, task_exec_map,
                completed_task_ids, total_tokens, total_cost, completed_count,
            )
            total_tokens = result["total_tokens"]
            total_cost = result["total_cost"]
            if result["completed"]:
                completed_count += 1

        return {"total_tokens": total_tokens, "total_cost": total_cost, "completed_count": completed_count}

    async def _execute_hierarchical(
        self, db, execution, tasks, agents_map, task_map,
        task_outputs, task_exec_map, completed_task_ids,
        total_tokens, total_cost, completed_count, total_tasks,
    ) -> Dict[str, Any]:
        """Hierarchical 模式：Manager 审查 Worker 输出。

        1. Worker agents 执行各自任务
        2. Manager agent 审查每个 Worker 的输出
        3. Manager 决定：通过 / 修改 / 重新分配
        """
        manager_keywords = ['manager', 'coordinator', 'orchestrator', 'reviewer', '管理者', '协调者', '审查者']
        manager_agent = None
        worker_agents = {}
        for aid, agent in agents_map.items():
            role_lower = (agent.role or '').lower()
            if any(kw in role_lower for kw in manager_keywords):
                manager_agent = agent
            else:
                worker_agents[aid] = agent

        if not manager_agent:
            # 无 Manager — 降级为 Sequential
            logger.warning("[HIERARCHICAL] No manager found, falling back to sequential")
            return await self._execute_sequential_fallback(
                db, execution, tasks, agents_map, task_map,
                task_outputs, task_exec_map, completed_task_ids,
                total_tokens, total_cost, completed_count, total_tasks,
            )

        max_review_rounds = 3
        for review_round in range(max_review_rounds):
            all_approved = True
            for task in tasks:
                if self._cancelled or self._paused:
                    break
                if task.id in completed_task_ids:
                    continue
                agent = worker_agents.get(task.agent_id, agents_map.get(task.agent_id))
                if not agent:
                    continue

                te = task_exec_map.get(task.id)
                result = await self._execute_task(
                    db, execution, task, agent, te, task_outputs, task_exec_map,
                    completed_task_ids, total_tokens, total_cost, completed_count,
                )
                total_tokens = result["total_tokens"]
                total_cost = result["total_cost"]
                worker_output = task_outputs.get(task.id, "")

                # Manager review
                if manager_agent and review_round < max_review_rounds - 1:
                    review_task = type('_ReviewTask', (), {
                        'id': f"{task.id}_review_{review_round}",
                        'name': f"审查: {task.name}",
                        'description': f"审查以下工作产出，判断是否通过（回复 APPROVED 或 REVISION_NEEDED: <原因>）：\n\n{worker_output[:2000]}",
                        'expected_output': 'APPROVED 或 REVISION_NEEDED: <具体修改建议>',
                        'agent_id': manager_agent.id,
                        'timeout_seconds': 120, 'max_retries': 0,
                    })()
                    review_te = None
                    review_result = await self._execute_task(
                        db, execution, review_task, manager_agent, review_te, task_outputs, task_exec_map,
                        completed_task_ids, total_tokens, total_cost, completed_count,
                    )
                    total_tokens = review_result["total_tokens"]
                    total_cost = review_result["total_cost"]
                    review_output = task_outputs.get(review_task.id, "")

                    if "APPROVED" in review_output.upper():
                        completed_task_ids.append(task.id)
                        completed_count += 1
                    else:
                        all_approved = False
                        task_outputs[task.id] = f"[需修改 - 审查意见] {review_output}\n\n[原始输出] {worker_output}"
                else:
                    completed_task_ids.append(task.id)
                    completed_count += 1

            if all_approved or review_round >= max_review_rounds - 1:
                break

        return {"total_tokens": total_tokens, "total_cost": total_cost, "completed_count": completed_count}

    async def _execute_evaluator_optimizer(
        self, db, execution, tasks, agents_map, task_map,
        task_outputs, task_exec_map, completed_task_ids,
        total_tokens, total_cost, completed_count, total_tasks,
    ) -> Dict[str, Any]:
        """Evaluator-Optimizer 模式：评估 + 优化质量循环。

        1. Generator agent 生成输出
        2. Evaluator agent 评分并给出改进建议
        3. Optimizer 根据反馈重新生成
        4. 循环直到质量达标
        """
        evaluator_keywords = ['evaluator', 'critic', 'reviewer', 'judge', '评估者', '审查者', '评判者']
        optimizer_keywords = ['optimizer', 'generator', 'writer', 'creator', '优化者', '生成者']

        evaluator_agent = None
        optimizer_agent = None
        for aid, agent in agents_map.items():
            role_lower = (agent.role or '').lower()
            if not evaluator_agent and any(kw in role_lower for kw in evaluator_keywords):
                evaluator_agent = agent
            elif not optimizer_agent and any(kw in role_lower for kw in optimizer_keywords):
                optimizer_agent = agent

        if not evaluator_agent or not optimizer_agent:
            logger.warning("[EVALUATOR_OPTIMIZER] Missing evaluator/optimizer, falling back to sequential")
            return await self._execute_sequential_fallback(
                db, execution, tasks, agents_map, task_map,
                task_outputs, task_exec_map, completed_task_ids,
                total_tokens, total_cost, completed_count, total_tasks,
            )

        max_optimize_rounds = 3
        quality_threshold = 7  # 1-10 scale

        for task in tasks:
            if self._cancelled or self._paused:
                break
            if task.id in completed_task_ids:
                continue

            for opt_round in range(max_optimize_rounds):
                # Generator produces output
                gen_task = type('_GenTask', (), {
                    'id': f"{task.id}_gen_{opt_round}",
                    'name': task.name,
                    'description': task.description + (
                        f"\n\n[优化反馈] {task_outputs.get(f'{task.id}_eval', '')}" if opt_round > 0 else ""
                    ),
                    'expected_output': task.expected_output,
                    'agent_id': optimizer_agent.id,
                    'timeout_seconds': getattr(task, 'timeout_seconds', 300),
                    'max_retries': getattr(task, 'max_retries', 1),
                })()
                gen_te = None
                gen_result = await self._execute_task(
                    db, execution, gen_task, optimizer_agent, gen_te, task_outputs, task_exec_map,
                    completed_task_ids, total_tokens, total_cost, completed_count,
                )
                total_tokens = gen_result["total_tokens"]
                total_cost = gen_result["total_cost"]
                gen_output = task_outputs.get(gen_task.id, "")

                # Evaluator scores
                eval_task = type('_EvalTask', (), {
                    'id': f"{task.id}_eval",
                    'name': f"评估: {task.name}",
                    'description': (
                        f"评估以下输出质量（1-10分），给出具体改进建议。\n"
                        f"回复格式：SCORE: <分数>\nFEEDBACK: <建议>\nDECISION: PASS/REVISE\n\n--- 待评估输出 ---\n{gen_output[:3000]}"
                    ),
                    'expected_output': 'SCORE: <1-10>\nFEEDBACK: <改进建议>\nDECISION: PASS|REVISE',
                    'agent_id': evaluator_agent.id,
                    'timeout_seconds': 120, 'max_retries': 0,
                })()
                eval_te = None
                eval_result = await self._execute_task(
                    db, execution, eval_task, evaluator_agent, eval_te, task_outputs, task_exec_map,
                    completed_task_ids, total_tokens, total_cost, completed_count,
                )
                total_tokens = eval_result["total_tokens"]
                total_cost = eval_result["total_cost"]
                eval_output = task_outputs.get(eval_task.id, "")

                # Check if passes quality threshold
                import re
                score_match = re.search(r'SCORE:\s*(\d+)', eval_output)
                score = int(score_match.group(1)) if score_match else 5

                if "PASS" in eval_output.upper() or score >= quality_threshold or opt_round >= max_optimize_rounds - 1:
                    task_outputs[task.id] = gen_output
                    completed_task_ids.append(task.id)
                    completed_count += 1
                    break
                else:
                    task_outputs[f"{task.id}_eval"] = eval_output

        return {"total_tokens": total_tokens, "total_cost": total_cost, "completed_count": completed_count}

    async def _execute_prompt_chain(
        self, db, execution, tasks, agents_map, task_map,
        task_outputs, task_exec_map, completed_task_ids,
        total_tokens, total_cost, completed_count, total_tasks,
    ) -> Dict[str, Any]:
        """Prompt Chain 模式：链式顺序，前一个任务的输出作为下一个任务的输入。"""
        logger.info(f"[PROMPT_CHAIN] Starting chain with {len(tasks)} task(s)")

        chain_context = ""
        for task in sorted(tasks, key=lambda t: t.order if hasattr(t, 'order') else 0):
            if self._cancelled or self._paused:
                break
            if task.id in completed_task_ids:
                continue

            agent = agents_map.get(task.agent_id)
            if not agent:
                continue

            if chain_context:
                task.description = f"{task.description}\n\n[链式上下文 — 上一步的输出]\n{chain_context[:2000]}"

            te = task_exec_map.get(task.id)
            result = await self._execute_task(
                db, execution, task, agent, te, task_outputs, task_exec_map,
                completed_task_ids, total_tokens, total_cost, completed_count,
            )
            total_tokens = result["total_tokens"]
            total_cost = result["total_cost"]
            if result["completed"]:
                completed_count += 1
                chain_context = task_outputs.get(task.id, chain_context)

        return {"total_tokens": total_tokens, "total_cost": total_cost, "completed_count": completed_count}

    async def _execute_router(
        self, db, execution, tasks, agents_map, task_map,
        task_outputs, task_exec_map, completed_task_ids,
        total_tokens, total_cost, completed_count, total_tasks,
    ) -> Dict[str, Any]:
        """Router 模式：Router agent 分析输入，路由到专门处理该类型任务的专家 agent。"""
        router_keywords = ['router', 'dispatcher', 'classifier', '路由', '分发', '分类']
        router_agent = None
        worker_agents = {}
        for aid, agent in agents_map.items():
            role_lower = (agent.role or '').lower()
            if not router_agent and any(kw in role_lower for kw in router_keywords):
                router_agent = agent
            else:
                worker_agents[aid] = agent

        if not router_agent:
            logger.warning("[ROUTER] No router found, using first agent as router")
            router_agent = list(agents_map.values())[0]

        for task in tasks:
            if self._cancelled or self._paused:
                break
            if task.id in completed_task_ids:
                continue

            # Route: Let router analyze and direct the task
            route_prompt = (
                f"分析以下任务并选择最合适的处理专家：\n\n{task.description[:1000]}\n\n"
                f"可用的专家及其角色：\n"
                + "\n".join(f"- {agent.name}: {agent.role}" for agent in worker_agents.values())
                + "\n\n回复格式：EXPERT: <专家名称>\nREASON: <选择理由>"
            )

            route_task = type('_RouteTask', (), {
                'id': f"{task.id}_route",
                'name': f"路由决策: {task.name}",
                'description': route_prompt,
                'expected_output': 'EXPERT: <name>\nREASON: <reason>',
                'agent_id': router_agent.id,
                'timeout_seconds': 120, 'max_retries': 0,
            })()
            route_te = None
            route_result = await self._execute_task(
                db, execution, route_task, router_agent, route_te, task_outputs, task_exec_map,
                completed_task_ids, total_tokens, total_cost, completed_count,
            )
            total_tokens = route_result["total_tokens"]
            total_cost = route_result["total_cost"]

            # Find the chosen expert
            route_output = task_outputs.get(route_task.id, "")
            chosen_agent = None
            for aid, agent in worker_agents.items():
                if agent.name.lower() in route_output.lower():
                    chosen_agent = agent
                    break
            if not chosen_agent:
                chosen_agent = agents_map.get(task.agent_id) or list(agents_map.values())[0]

            # Execute with chosen expert
            te = task_exec_map.get(task.id)
            result = await self._execute_task(
                db, execution, task, chosen_agent, te, task_outputs, task_exec_map,
                completed_task_ids, total_tokens, total_cost, completed_count,
            )
            total_tokens = result["total_tokens"]
            total_cost = result["total_cost"]
            if result["completed"]:
                completed_count += 1

        return {"total_tokens": total_tokens, "total_cost": total_cost, "completed_count": completed_count}

    async def _execute_sequential_fallback(
        self, db, execution, tasks, agents_map, task_map,
        task_outputs, task_exec_map, completed_task_ids,
        total_tokens, total_cost, completed_count, total_tasks,
    ) -> Dict[str, Any]:
        """退路：顺序执行（当特殊模式的 Agent 不可用时）。"""
        levels = self._topological_levels(tasks, task_map)
        for level in levels:
            for task_id in level:
                if self._cancelled or self._paused:
                    break
                if task_id in completed_task_ids:
                    continue
                task = task_map.get(task_id)
                if not task:
                    continue
                agent = agents_map.get(task.agent_id)
                if not agent:
                    continue
                te = task_exec_map.get(task.id)
                result = await self._execute_task(
                    db, execution, task, agent, te, task_outputs, task_exec_map,
                    completed_task_ids, total_tokens, total_cost, completed_count,
                )
                total_tokens = result["total_tokens"]
                total_cost = result["total_cost"]
                if result["completed"]:
                    completed_count += 1
        return {"total_tokens": total_tokens, "total_cost": total_cost, "completed_count": completed_count}


def cancel_execution_engine(execution_id: str) -> bool:
    engine = _running_engines.get(execution_id)
    if engine:
        engine.cancel()
        return True
    return False


def pause_execution_engine(execution_id: str) -> bool:
    """向正在运行的引擎发送暂停信号（内存中的快速路径）"""
    engine = _running_engines.get(execution_id)
    if engine:
        engine.pause()
        return True
    return False
